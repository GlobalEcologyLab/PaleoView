# Python modules
from os import chdir, getcwd, system, path, listdir
import filecmp
import string

# Python extension modules (requires extension installation)
import numpy as np
import pandas as pd
from netCDF4 import Dataset
from docx import Document

# Tool library module
from PaleoclimateToolDataFileHelper import PaleoclimateToolDataFileHelper

## Main program

# Create the Paleoclimate Tool Data File Helper
data_helper = PaleoclimateToolDataFileHelper()

# Set the climate data, mask, and bias correction directories
chdir(path.join(getcwd(), '..'))
current_tool_version_directory = getcwd()
chdir(path.join(getcwd(), '..'))
tool_directory = getcwd()
data_helper.setClimateDataSource('local')
#data_helper.setClimateDataDirectory(path.join(tool_directory, 'Climate Data'))
data_helper.setClimateDataDirectory(r'C:/Users/shaythorne/Desktop/PaleoView/Tests/NetCdfClimateData')
data_helper.getCurrentNetCdfDataIntervals()
data_helper.setRegionMaskDirectory(path.join(current_tool_version_directory, 'Map Data'))
data_helper.setBiasCorrectionDirectory(path.join(current_tool_version_directory, 'Bias Corrections'))
data_helper.setFileGenerationDirectory(path.join(current_tool_version_directory, 'Test', 'DataFileTests'))

### TEST loadClimateDataGrids
##print 'TEST loadClimateDataGrids:'
##directory = data_helper.climate_data_directory['path']
##
##print '  Test 1: mean temperature, 1989 AD, months A-E'
##data_grids1 = data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=1989, month_indices=range(5))
##data_grids = [] # expected
##for month_code in ['A', 'B', 'C', 'D', 'E'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'T', 'Trace21_2.5x2.5_1989AD.1T'+month_code+'.txt')))  
##print '    Pass =', (np.array(data_grids) == data_grids1.round(3)).all()
##
##print '  Test 2: minimum temperature, 10 BP (1940), months A-L (all)'
##data_grids2 = data_helper.loadClimateDataGrids(parameter='minimum_temperature', year_ad=1940, month_indices=range(12))
##data_grids = [] # expected
##for month_code in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'Tmin', 'Trace21_2.5x2.5_10BP.1I'+month_code+'.txt')))  
##print '    Pass =', (np.array(data_grids) == data_grids2.round(3)).all()
##
##print '  Test 3: maximum temperature, 1989 AD, months A,B,C,K,L'
##data_grids3 = data_helper.loadClimateDataGrids(parameter='maximum_temperature', year_ad=1989, month_indices=[0,1,2,10,11])
##data_grids = [] # expected
##for month_code in ['K', 'L'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'Tmax', 'Trace21_2.5x2.5_1988AD.1A'+month_code+'.txt')))  
##for month_code in ['A', 'B', 'C'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'Tmax', 'Trace21_2.5x2.5_1989AD.1A'+month_code+'.txt')))  
##print '    Pass =', (np.array(data_grids) == data_grids3.round(3)).all()
##
##print '  Test 4: relative_humidity, 1951 AD, months A,B,C,K,L'
##data_grids4 = data_helper.loadClimateDataGrids(parameter='relative_humidity', year_ad=1951, month_indices=[0,1,2,10,11])
##data_grids = [] # expected
##for month_code in ['K', 'L'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'H', 'Trace21_2.5x2.5_0BP.1H'+month_code+'.txt')))  
##for month_code in ['A', 'B', 'C'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'H', 'Trace21_2.5x2.5_1951AD.1H'+month_code+'.txt')))  
##print '    Pass =', (np.array(data_grids) == data_grids4.round(6)).all()
##
##print '  Test 5: specific humidity, 1989 AD, month G'
##data_grids5 = data_helper.loadClimateDataGrids(parameter='specific_humidity', year_ad=1989, month_indices=[6])
##data_grids = [np.genfromtxt(path.join(directory, 'Q', 'Trace21_2.5x2.5_1989AD.1SG.txt'))] # expected
##print '    Pass =', (np.array(data_grids) == data_grids5.round(3)).all()
##
##print '  Test 6: precipitation, 1989 AD, month G'
##data_grids6 = data_helper.loadClimateDataGrids(parameter='precipitation', year_ad=1989, month_indices=[6])
##data_grids = [np.genfromtxt(path.join(directory, 'P', 'Trace21_2.5x2.5_1989AD.1PG.txt'))] # expected
##print '    Pass =', (np.array(data_grids) == data_grids6.round(3)).all()
##
### TEST loadNongriddedDataFrame
##print 'TEST loadNongriddedDataFrame:'
##full_soi_dataframe = pd.read_csv(path.join(directory, 'South_Oscillation_Index.txt'), delim_whitespace=True)
##
##print '  Test 1: soi, 10-5 BP (1939-45)'
##data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation',
##                                    parameter_code='soi',
##                                    period_ad_from=1940,
##                                    period_ad_until=1945)
##data_frame = full_soi_dataframe[((full_soi_dataframe['BP/AD'] == 'BP') & (full_soi_dataframe['Yrs'] >= 5) & (full_soi_dataframe['Yrs'] <= 11))]
##print '    Pass =', (data_frame.get_values() == data_helper.nongridded_data_frame.get_values()).all()
##
##print '  Test 2: soi, 5 BP - 1955 AD (1944-1955)'
##data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation',
##                                    parameter_code='soi',
##                                    period_ad_from=1945,
##                                    period_ad_until=1955)
##data_frame = full_soi_dataframe[( ((full_soi_dataframe['BP/AD'] == 'BP') & (full_soi_dataframe['Yrs'] >= 0) & (full_soi_dataframe['Yrs'] <= 6)) |
##                                  ((full_soi_dataframe['BP/AD'] == 'AD') & (full_soi_dataframe['Yrs'] >= 1951) & (full_soi_dataframe['Yrs'] <= 1955)) )]
##print '    Pass =', (data_frame.get_values() == data_helper.nongridded_data_frame.get_values()).all()
##
##print '  Test 3: soi, 1979-89 AD'
##data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation',
##                                    parameter_code='soi',
##                                    period_ad_from=1980,
##                                    period_ad_until=1989)
##data_frame = full_soi_dataframe[((full_soi_dataframe['BP/AD'] == 'AD') & (full_soi_dataframe['Yrs'] >= 1979) & (full_soi_dataframe['Yrs'] <= 1989))]
##print '    Pass =', (data_frame.get_values() == data_helper.nongridded_data_frame.get_values()).all()
##
##print '  Test 4: soi, 5 BP - 1955 AD (1944-1955) and delta period 100-95BP (1849-1855)'
##data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation',
##                                    parameter_code='soi',
##                                    period_ad_from=1945,
##                                    period_ad_until=1955,
##                                    delta_interval={ 'ad_from' : 1850, 'ad_until' : 1855 })
##data_frame = full_soi_dataframe[( ((full_soi_dataframe['BP/AD'] == 'BP') & (full_soi_dataframe['Yrs'] >= 0) & (full_soi_dataframe['Yrs'] <= 6)) |
##                                  ((full_soi_dataframe['BP/AD'] == 'AD') & (full_soi_dataframe['Yrs'] >= 1951) & (full_soi_dataframe['Yrs'] <= 1955)) |
##                                  ((full_soi_dataframe['BP/AD'] == 'BP') & (full_soi_dataframe['Yrs'] >= 95) & (full_soi_dataframe['Yrs'] <= 101)) )]
##print '    Pass =', (data_frame.get_values() == data_helper.nongridded_data_frame.get_values()).all()
##
### TEST loadNongriddedClimateData
##print 'TEST loadNongriddedClimateData:'
##data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation',
##                                    parameter_code='soi',
##                                    period_ad_from=1940,
##                                    period_ad_until=1989)
##data_frame = data_helper.nongridded_data_frame
##
##print '  Test 1: soi, 1989 AD, months A-E'
##nongridded_data1 = data_helper.loadNongriddedClimateData(parameter_group='southern-oscillation',
##                                                        parameter='soi',
##                                                        year_ad=1989,
##                                                        month_indices=range(5))
##values = data_frame[(data_frame['BP/AD'] == 'AD') & (data_frame['Yrs'] == 1989) & (data_frame['Months'] <= 'E')]['SOI'].get_values()
##print '    Pass =', (values == nongridded_data1).all()
##
##print '  Test 2: soi, 10 BP (1940), months A-L (all)'
##nongridded_data2 = data_helper.loadNongriddedClimateData(parameter_group='southern-oscillation',
##                                                        parameter='soi',
##                                                        year_ad=1940,
##                                                        month_indices=range(12))
##values = data_frame[(data_frame['BP/AD'] == 'BP') & (data_frame['Yrs'] == 10)]['SOI'].get_values()
##print '    Pass =', (values == nongridded_data2).all()
##
##print '  Test 3: soi, 1989 AD, months A,B,C,K,L'
##nongridded_data3 = data_helper.loadNongriddedClimateData(parameter_group='southern-oscillation',
##                                                        parameter='soi',
##                                                        year_ad=1989,
##                                                        month_indices=[0,1,2,10,11])
##values_88 = data_frame[(data_frame['BP/AD'] == 'AD') & (data_frame['Yrs'] == 1988) & (data_frame['Months'] >= 'K')]['SOI'].get_values()
##values_89 = data_frame[(data_frame['BP/AD'] == 'AD') & (data_frame['Yrs'] == 1989) & (data_frame['Months'] <= 'C')]['SOI'].get_values()
##values = values_88.tolist()
##values.extend(values_89)
##print '    Pass =', (np.array(values) == nongridded_data3).all()
##
### TEST generateParameterDataInterval
##print 'TEST generateParameterDataInterval:'
##
##print '  Test 1: mean temperature, 1980-9 AD, months A,B, region full, grid, averaged'
##data_interval1 = data_helper.generateParameterDataInterval(parameter_group_code='temperature',
##                                                           parameter_code='mean-temperature',
##                                                           interval_ad_from=1980,
##                                                           interval_ad_until=1989,
##                                                           month_indices=[0,1])
##data_grids = []
##for year_ad in range(1980, 1990) :
##    data_grids.append(data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=year_ad, month_indices=[0,1]))
##data_interval = np.array(data_grids).mean(1).mean(0)
##print '    Pass =', (data_interval == data_interval1).all()
##
##print '  Test 2: diurnal temperature range, 1980-9 AD, months A,B, region full, grid, averaged'
##data_interval2 = data_helper.generateParameterDataInterval(parameter_group_code='temperature',
##                                                           parameter_code='diurnal-temperature-range',
##                                                           interval_ad_from=1980,
##                                                           interval_ad_until=1989,
##                                                           month_indices=[0,1])
##data_grids_for_min = []
##data_grids_for_max = []
##for year_ad in range(1980, 1990) :
##    data_grids_for_min.append(data_helper.loadClimateDataGrids(parameter='minimum_temperature', year_ad=year_ad, month_indices=[0,1]))
##    data_grids_for_max.append(data_helper.loadClimateDataGrids(parameter='maximum_temperature', year_ad=year_ad, month_indices=[0,1]))
##data_interval = (np.array(data_grids_for_max) - np.array(data_grids_for_min)).mean(1).mean(0)
##print '    Pass =', (data_interval == data_interval2).all()
##
##print '  Test 3: temperature seasonality, 1980-9 AD, months A-L region full, grid, averaged'
##data_interval3 = data_helper.generateParameterDataInterval(parameter_group_code='temperature',
##                                                           parameter_code='temperature-seasonality',
##                                                           interval_ad_from=1980,
##                                                           interval_ad_until=1989,
##                                                           month_indices=range(12))
##data_grids = []
##for year_ad in range(1980, 1990) :
##    data_grids.append(data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=year_ad, month_indices=range(12)))
##data_interval = (np.array(data_grids).std(1)*100).mean(0)
##print '    Pass =', (data_interval == data_interval3).all()
##
##print '  Test 4: precipitation seasonality, 1980-9 AD, months A-L region full, grid, averaged'
##data_interval4 = data_helper.generateParameterDataInterval(parameter_group_code='precipitation',
##                                                           parameter_code='precipitation-seasonality',
##                                                           interval_ad_from=1980,
##                                                           interval_ad_until=1989,
##                                                           month_indices=range(12))
##data_grids = []
##for year_ad in range(1980, 1990) :
##    data_grids.append(data_helper.loadClimateDataGrids(parameter='precipitation', year_ad=year_ad, month_indices=range(12)))
##data_interval = (np.array(data_grids).std(1)/np.array(data_grids).mean(1)).mean(0)
##print '    Pass =', (data_interval == data_interval4).all()
##
##print '  Test 5: annual temperature range, 1980-9 AD, months A-L region full, grid, averaged'
##data_interval5 = data_helper.generateParameterDataInterval(parameter_group_code='temperature',
##                                                           parameter_code='annual-temperature-range',
##                                                           interval_ad_from=1980,
##                                                           interval_ad_until=1989,
##                                                           month_indices=range(12))
##data_grids = []
##for year_ad in range(1980, 1990) :
##    data_grids.append(data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=year_ad, month_indices=range(12)))
##data_interval = (np.array(data_grids).max(1)-np.array(data_grids).min(1)).mean(0)
##print '    Pass =', (data_interval == data_interval5).all()
##
##print '  Test 6: isothermality, 1980-9 AD, months A-L region full, grid, averaged'
##data_interval6 = data_helper.generateParameterDataInterval(parameter_group_code='temperature',
##                                                           parameter_code='isothermality',
##                                                           interval_ad_from=1980,
##                                                           interval_ad_until=1989,
##                                                           month_indices=range(12))
##data_grids_for_mean = []
##data_grids_for_min = []
##data_grids_for_max = []
##for year_ad in range(1980, 1990) :
##    data_grids_for_mean.append(data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=year_ad, month_indices=range(12)))
##    data_grids_for_min.append(data_helper.loadClimateDataGrids(parameter='minimum_temperature', year_ad=year_ad, month_indices=range(12)))
##    data_grids_for_max.append(data_helper.loadClimateDataGrids(parameter='maximum_temperature', year_ad=year_ad, month_indices=range(12)))
##data_interval = ((np.array(data_grids_for_max) - np.array(data_grids_for_min)).mean(1)/(np.array(data_grids_for_mean).max(1)-np.array(data_grids_for_mean).min(1))*100).mean(0)
##print '    Pass =', (data_interval == data_interval6).all()
##
##print '  Test 7: soi, 1980-9 AD, months A,B, region full, no grid, average'
##data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation', parameter_code='soi', period_ad_from=1980, period_ad_until=1989)
##data_interval7 = data_helper.generateParameterDataInterval(parameter_group_code='southern-oscillation',
##                                                           parameter_code='soi',
##                                                           interval_ad_from=1980,
##                                                           interval_ad_until=1989,
##                                                           month_indices=[0,1])
##data_frame = data_helper.nongridded_data_frame
##values = []
##for month in ['A', 'B'] :
##    values.append(data_frame[(data_frame['BP/AD'] == 'AD') & (data_frame['Yrs'] >= 1980) & (data_frame['Months'] == month)]['SOI'].get_values())
##print '    Pass =', (np.array(values).mean(0).mean(0) == data_interval7)
##
##print '  Test 8: soi, 1980-9 AD, months A-L, region full, no grid'
##data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation', parameter_code='soi', period_ad_from=1980, period_ad_until=1989)
##data_interval8 = data_helper.generateParameterDataInterval(parameter_group_code='southern-oscillation',
##                                                           parameter_code='soi',
##                                                           interval_ad_from=1980,
##                                                           interval_ad_until=1989,
##                                                           month_indices=range(12))
##data_frame = data_helper.nongridded_data_frame
##values = []
##for month in data_helper.month_codes :
##    values.append(data_frame[(data_frame['BP/AD'] == 'AD') & (data_frame['Yrs'] >= 1980) & (data_frame['Months'] == month)]['SOI'].get_values())
##print '    Pass =', (np.array(values).mean(0).mean(0) == data_interval8)
##
### TEST calculateGridRegionStatistics
##print 'TEST calculateGridRegionStatistics:'
##data_grid = data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=1989, month_indices=[0])[0]
##region_mask = np.zeros((data_helper.grid_height, data_helper.grid_width))
##region_mask[0:10,0:10] = 1
##data_grid[0,0] = 0.0*data_grid[0,0]/0.0
##data_grid[1,0] = -1.0*data_grid[1,0]/0.0
##data_grid[0,1] = 1.0*data_grid[0,1]/0.0
##grid_region_statistics = data_helper.calculateGridRegionStatistics([data_grid], [region_mask])
##region_data = np.ma.masked_array(data_grid[0:10,0:10], mask=((np.isfinite(data_grid[0:10,0:10]) - 1)*-1))
##cos_lat = np.meshgrid(np.ones(10), np.cos(np.arange(88.75, 66.2499, -2.5)*np.pi/180.0))[1]
##area_average = (region_data*cos_lat).sum()/cos_lat.sum()
##stats = { 'minimum' : [region_data.min()], 'maximum' : [region_data.max()], 'percentile_5th' : [np.percentile(region_data, 5)],
##          'percentile_25th' : [np.percentile(region_data, 25)], 'percentile_50th' : [np.percentile(region_data, 50)],
##          'percentile_75th' : [np.percentile(region_data, 75)], 'percentile_95th' : [np.percentile(region_data, 95)],
##          'grid_mean' : [region_data.mean()], 'grid_stdev' : [region_data.std()],
##          'area_mean' : [area_average], 'area_stdev' : [np.sqrt(((area_average - region_data)**2*cos_lat).sum()/cos_lat.sum())] }
##print '  Pass =', grid_region_statistics == stats
##
### TEST generateParameterData
##print 'TEST generateParameterData:'
##
##print '  Test 1: minimum temperature, 1950-1980, interval steps 10 yr, size 10 yr, months B, region full, grid, averaged'
##parameter_data1 = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                    parameter_code='minimum-temperature',
##                                                    period_ad_from=1950,
##                                                    period_ad_until=1980,
##                                                    interval_step=10,
##                                                    interval_size=10,
##                                                    month_indices=[1],
##                                                    region_mask=1,
##                                                    generate_grids=True,
##                                                    all_months=False)
##data_grids = []
##for from_year_ad in range(1950, 1981, 10) :
##    interval_grids = []
##    for year_ad in range(from_year_ad-5, from_year_ad+5) :
##        interval_grids.append(data_helper.loadClimateDataGrids(parameter='minimum_temperature', year_ad=year_ad, month_indices=[1]))
##    data_grids.append(np.array(interval_grids).mean(1).mean(0))
##print '    Pass =', (np.array(data_grids) == np.array(parameter_data1)).all()
##
##region_mask = np.zeros((data_helper.grid_height, data_helper.grid_width))
##region_mask[0:4,0:4] = 1
##
##print '  Test 2: mean temperature, 1950-1980, interval steps 10 yr, size 10 yr, months A,B, region TL4x4, non-grid, averaged'
##parameter_data2 = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                    parameter_code='mean-temperature',
##                                                    period_ad_from=1950,
##                                                    period_ad_until=1980,
##                                                    interval_step=10,
##                                                    interval_size=10,
##                                                    month_indices=[0,1],
##                                                    region_mask=region_mask,
##                                                    generate_grids=False,
##                                                    all_months=False)
##data_grids = []
##region_masks = []
##for from_year_ad in range(1950, 1981, 10) :
##    interval_grids = []
##    for year_ad in range(from_year_ad-5, from_year_ad+5) :
##        interval_grids.append(data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=year_ad, month_indices=[0,1]))
##    data_grids.append(np.array(interval_grids).mean(1).mean(0))
##    region_masks.append(region_mask)
##stats_data = data_helper.calculateGridRegionStatistics(data_grids, region_masks)
##print '    Pass =', (stats_data == parameter_data2)
##
##print '  Test 3: mean temperature, 1950-1980, interval steps 10 yr, size 5 yr, months A,B, region full, grid, averaged'
##parameter_data3 = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                    parameter_code='mean-temperature',
##                                                    period_ad_from=1950,
##                                                    period_ad_until=1980,
##                                                    interval_step=10,
##                                                    interval_size=5,
##                                                    month_indices=[0,1],
##                                                    region_mask=1,
##                                                    generate_grids=True,
##                                                    all_months=False)
##data_grids = []
##for from_year_ad in range(1950, 1981, 10) :
##    interval_grids = []
##    for year_ad in range(from_year_ad-2, from_year_ad+3) :
##        interval_grids.append(data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=year_ad, month_indices=[0,1]))
##    data_grids.append(np.array(interval_grids).mean(1).mean(0))
##print '    Pass =', (np.array(data_grids) == np.array(parameter_data3)).all()
##
##region_mask = np.zeros((data_helper.grid_height, data_helper.grid_width))
##region_mask[0:4,0:4] = 1
##
##print '  Test 4: minimum temperature, 1950-1980, interval steps 10 yr, size 10 yr, all months, region TL4x4, non-grid'
##parameter_data4 = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                    parameter_code='minimum-temperature',
##                                                    period_ad_from=1950,
##                                                    period_ad_until=1980,
##                                                    interval_step=10,
##                                                    interval_size=10,
##                                                    month_indices=range(12),
##                                                    region_mask=region_mask,
##                                                    generate_grids=False,
##                                                    all_months=True)
##stats_data = []
##for month_index in range(12) :
##    data_grids_month = []
##    region_masks = []
##    for from_year_ad in range(1950, 1981, 10) :
##        interval_grids = []
##        for year_ad in range(from_year_ad-5, from_year_ad+5) :
##            interval_grids.append(data_helper.loadClimateDataGrids(parameter='minimum_temperature', year_ad=year_ad, month_indices=[month_index]))
##        data_grids_month.append(np.array(interval_grids).mean(1).mean(0))
##        region_masks.append(region_mask)
##    stats_data.append(data_helper.calculateGridRegionStatistics(data_grids_month, region_masks))
##print '    Pass =', (stats_data == parameter_data4)
##
##print '  Test 5: soi, 1930-1980, interval steps 10 yr, size 10 yr, months A,B, no grid, averaged'
##parameter_data5 = data_helper.generateParameterData(parameter_group_code='southern-oscillation',
##                                                    parameter_code='soi',
##                                                    period_ad_from=1930,
##                                                    period_ad_until=1980,
##                                                    interval_step=10,
##                                                    interval_size=10,
##                                                    month_indices=[0,1],
##                                                    region_mask=1,
##                                                    generate_grids=False,
##                                                    all_months=False)
##data_frame = data_helper.nongridded_data_frame
##values = []
##for month in data_helper.month_codes[0:2] :
##    bp_values = data_frame[(data_frame['BP/AD'] == 'BP') & (data_frame['Yrs'] <= 25) & (data_frame['Months'] == month)]['SOI'].get_values()
##    ad_values = data_frame[(data_frame['BP/AD'] == 'AD') & (data_frame['Yrs'] >= 1951) & (data_frame['Yrs'] <= 1984) & (data_frame['Months'] == month)]['SOI'].get_values()
##    month_values = bp_values.tolist()
##    month_values.reverse()
##    month_values.extend(ad_values.tolist())
##    month_values = np.array(month_values).reshape((6, 10))
##    values.append(month_values)
##values = np.array(values).mean(0).mean(1)
##print '    Pass =', (values == np.array(parameter_data5)).all()
##
##print '  Test 6: soi, 1960-1980, interval steps 10 yr, size 5 yr, no grid, all months'
##parameter_data6 = data_helper.generateParameterData(parameter_group_code='southern-oscillation',
##                                                    parameter_code='soi',
##                                                    period_ad_from=1960,
##                                                    period_ad_until=1980,
##                                                    interval_step=10,
##                                                    interval_size=5,
##                                                    month_indices=range(12),
##                                                    region_mask=1,
##                                                    generate_grids=False,
##                                                    all_months=True)
##data_frame = data_helper.nongridded_data_frame
##values = []
##for month_index in range(12) :
##    values_month = []
##    for from_year_ad in range(1960, 1981, 10) :
##        data = data_frame[(data_frame['BP/AD'] == 'AD') & (data_frame['Yrs'] >= (from_year_ad-2)) & (data_frame['Yrs'] <= (from_year_ad+2)) & (data_frame['Months'] == data_helper.month_codes[month_index])]['SOI'].get_values()
##        values_month.append(data)
##    values.append(np.array(values_month).mean(1))
##print '    Pass =', (np.array(values) == np.array(parameter_data6)).all()
##
##print '  Test 7: delta mean temperature, 1950-1980 relative to 1940, interval steps 10 yr, size 10 yr, months A,B, region full, grid, averaged'
##parameter_data7 = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                    parameter_code='mean-temperature',
##                                                    period_ad_from=1950,
##                                                    period_ad_until=1980,
##                                                    delta_ref_period_ad=1940,
##                                                    interval_step=10,
##                                                    interval_size=10,
##                                                    month_indices=[0,1],
##                                                    region_mask=1,
##                                                    generate_grids=True,
##                                                    all_months=False)
##data_grids = []
##for from_year_ad in range(1940, 1981, 10) :
##    interval_grids = []
##    for year_ad in range(from_year_ad-5, from_year_ad+5) :
##        interval_grids.append(data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=year_ad, month_indices=[0,1]))
##    data_grids.append(np.array(interval_grids).mean(1).mean(0))
##delta_grids = []
##for i in range(4) :
##    delta_grids.append(data_grids[i+1] - data_grids[0])
##print '    Pass =', (np.array(delta_grids) == np.array(parameter_data7)).all()
##
##print '  Test 8: delta mean temperature, 1950-1980 relative to 1940, interval steps 10 yr, size 10 yr, months A,B, region TL4x4, non-grid, averaged'
##region_mask = np.zeros((data_helper.grid_height, data_helper.grid_width))
##region_mask[0:4,0:4] = 1
##parameter_data8 = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                    parameter_code='mean-temperature',
##                                                    period_ad_from=1950,
##                                                    period_ad_until=1980,
##                                                    delta_ref_period_ad=1940,
##                                                    interval_step=10,
##                                                    interval_size=10,
##                                                    month_indices=[0,1],
##                                                    region_mask=region_mask,
##                                                    generate_grids=False,
##                                                    all_months=False)
##region_masks = [region_mask, region_mask, region_mask, region_mask]
##stats_data = data_helper.calculateGridRegionStatistics(delta_grids, region_masks)
##print '    Pass =', (stats_data == parameter_data8)
##
##print '  Test 9: delta minimum temperature, 1950-1980 relative to 1940, interval steps 10 yr, size 10 yr, all months, region TL4x4, non-grid'
##parameter_data9 = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                    parameter_code='minimum-temperature',
##                                                    period_ad_from=1950,
##                                                    period_ad_until=1980,
##                                                    delta_ref_period_ad=1940,
##                                                    interval_step=10,
##                                                    interval_size=10,
##                                                    month_indices=range(12),
##                                                    region_mask=region_mask,
##                                                    generate_grids=False,
##                                                    all_months=True)
##stats_data = []
##for month_index in range(12) :
##    data_grids_month = []
##    for from_year_ad in range(1940, 1981, 10) :
##        interval_grids = []
##        for year_ad in range(from_year_ad-5, from_year_ad+5) :
##            interval_grids.append(data_helper.loadClimateDataGrids(parameter='minimum_temperature', year_ad=year_ad, month_indices=[month_index]))
##        data_grids_month.append(np.array(interval_grids).mean(1).mean(0))
##    delta_grids_month = []
##    for i in range(4) :
##        delta_grids_month.append(data_grids_month[i+1] - data_grids_month[0])
##    stats_data.append(data_helper.calculateGridRegionStatistics(delta_grids_month, region_masks))
##print '    Pass =', (stats_data == parameter_data9)
##
##print '  Test 10: % delta mean precipitation, 1950-1980 relative to 1940, interval steps 10 yr, size 10 yr, months A,B, region full, grid, averaged'
##parameter_data10 = data_helper.generateParameterData(parameter_group_code='precipitation',
##                                                    parameter_code='mean-precipitation',
##                                                    period_ad_from=1950,
##                                                    period_ad_until=1980,
##                                                    delta_ref_period_ad=1940,
##                                                    delta_as_percent=True,
##                                                    interval_step=10,
##                                                    interval_size=10,
##                                                    month_indices=[0,1],
##                                                    region_mask=1,
##                                                    generate_grids=True,
##                                                    all_months=False)
##data_grids = []
##for from_year_ad in range(1940, 1981, 10) :
##    interval_grids = []
##    for year_ad in range(from_year_ad-5, from_year_ad+5) :
##        interval_grids.append(data_helper.loadClimateDataGrids(parameter='precipitation', year_ad=year_ad, month_indices=[0,1]))
##    data_grids.append(np.array(interval_grids).mean(1).mean(0))
##delta_grids = []
##for i in range(4) :
##    delta_grids.append(100*(data_grids[i+1] - data_grids[0])/data_grids[0])
##print '    Pass =', (np.array(delta_grids) == np.array(parameter_data10)).all()
##
##print '  Test 11: delta soi, 1950-1980 relative to 1940, interval steps 10 yr, size 10 yr, months A,B, no grid, averaged'
##parameter_data11 = data_helper.generateParameterData(parameter_group_code='southern-oscillation',
##                                                    parameter_code='soi',
##                                                    period_ad_from=1950,
##                                                    period_ad_until=1980,
##                                                    delta_ref_period_ad=1940,
##                                                    interval_step=10,
##                                                    interval_size=10,
##                                                    month_indices=[0,1],
##                                                    region_mask=1,
##                                                    generate_grids=False,
##                                                    all_months=False)
##data_frame = data_helper.nongridded_data_frame
##values = []
##delta_values = []
##for month in data_helper.month_codes[0:2] :
##    delta_bp_values = data_frame[(data_frame['BP/AD'] == 'BP') & (data_frame['Yrs'] <= 15) & (data_frame['Yrs'] >= 6) & (data_frame['Months'] == month)]['SOI'].get_values()
##    bp_values = data_frame[(data_frame['BP/AD'] == 'BP') & (data_frame['Yrs'] <= 5) & (data_frame['Months'] == month)]['SOI'].get_values()
##    ad_values = data_frame[(data_frame['BP/AD'] == 'AD') & (data_frame['Yrs'] >= 1951) & (data_frame['Yrs'] <= 1984) & (data_frame['Months'] == month)]['SOI'].get_values()
##    month_values = bp_values.tolist()
##    month_values.reverse()
##    month_values.extend(ad_values.tolist())
##    month_values = np.array(month_values).reshape((4, 10))
##    values.append(month_values)
##    delta_month_values = delta_bp_values.tolist()
##    delta_month_values.reverse()
##    delta_month_values = np.array(delta_month_values).reshape((1, 10))
##    delta_values.append(delta_month_values)
##values = np.array(values).mean(0).mean(1)
##delta_values = np.array(delta_values).mean(0).mean(1)
##print '    Pass =', ((values - delta_values) == np.array(parameter_data11)).all()
##
##print '  Test 12: % delta soi, 1950-1980 relative to 1940, interval steps 10 yr, size 10 yr, all months'
##parameter_data12 = data_helper.generateParameterData(parameter_group_code='southern-oscillation',
##                                                    parameter_code='soi',
##                                                    period_ad_from=1960,
##                                                    period_ad_until=1980,
##                                                    delta_ref_period_ad=1940,
##                                                    delta_as_percent=True,
##                                                    interval_step=10,
##                                                    interval_size=5,
##                                                    month_indices=range(12),
##                                                    region_mask=1,
##                                                    generate_grids=False,
##                                                    all_months=True)
##data_frame = data_helper.nongridded_data_frame
##values = []
##delta_ref_values = []
##for month_index in range(12) :
##    values_month = []
##    for from_year_ad in range(1960, 1981, 10) :
##        data = data_frame[(data_frame['BP/AD'] == 'AD') & (data_frame['Yrs'] >= (from_year_ad-2)) & (data_frame['Yrs'] <= (from_year_ad+2)) & (data_frame['Months'] == data_helper.month_codes[month_index])]['SOI'].get_values()
##        values_month.append(data)
##    values.append(np.array(values_month).mean(1))
##    delta_ref_values_month = data_frame[(data_frame['BP/AD'] == 'BP') & (data_frame['Yrs'] >= 8) & (data_frame['Yrs'] <= 12) & (data_frame['Months'] == data_helper.month_codes[month_index])]['SOI'].get_values()
##    delta_ref_values.append(np.array(delta_ref_values_month).mean(0))
##delta_ref_values = np.meshgrid(np.ones(3), np.array(delta_ref_values))[1]
##delta_values_as_percent = 100.0*(np.array(values) - delta_ref_values)/delta_ref_values
##print '    Pass =', (delta_values_as_percent.round(8) == np.array(parameter_data12).round(8)).all()
##
### TEST calculateDeltaValues (delta as percentage cases)
##print 'TEST calculateDeltaValues (delta as percentage cases):'
##
##print '  Test 1: grid data with divide by zero cases'
##parameter_data = [np.array([[0.0, 2.0, -1.0, 1.0, 0.0]])]
##delta_ref_data = np.array([[1.0, 1.0, 0.0, 0.0, 0.0]])
##delta_data1 = data_helper.calculateDeltaValues(parameter_data, delta_ref_data, delta_as_percent=True, grid_data=True)
##print '    Pass =', (delta_data1[0] == np.array([[-100.0, 100.0, -100.0, 100.0, 0.0]])).all()
##
##print '  Test 2: non-grid data with divide by zero cases'
##parameter_data = [-1.0, 1.0, 0.0]
##delta_ref_data = 0.0
##delta_data2 = data_helper.calculateDeltaValues(parameter_data, delta_ref_data, delta_as_percent=True, grid_data=False)
##print '    Pass =', (np.array(delta_data2) == np.array([-100.0, 100.0, 0.0])).all()
##
### TEST loadRegionMask
##print 'TEST loadRegionMask:'
##
##print '  Test 1: fixed mask'
##print '    Pass =', data_helper.loadRegionMask('globe').all()
##
##print '  Test 2: time dependent mask'
##region_mask_dict = data_helper.loadRegionMask('land-0-21KBP', time_dependent=True)
##keys = region_mask_dict.keys()
##keys.sort()
##data_helper.setRegionMaskDirectory(path.join(current_tool_version_directory, 'Map Data', 'land-0-21KBP'))
##region_mask = data_helper.loadRegionMask('land-11500BP', time_dependent=False)
##data_helper.setRegionMaskDirectory(path.join(current_tool_version_directory, 'Map Data'))
##print '    Pass =', len(region_mask_dict) == 211 and (keys == np.arange(0, 21001, 100)).all() and (region_mask_dict[11500] == region_mask).all()
##
### TEST loadBiasCorrectionDataGrids
##print 'TEST loadBiasCorrectionDataGrids:'
##directory = data_helper.bias_correction_directory['path']
##
##print '  Test 1: mean temperature, months A-E'
##bias_data_grids1 = data_helper.loadBiasCorrectionDataGrids(parameter='mean_temperature', month_indices=range(5))
##bias_data_grids = [] # expected
##for month_code in ['A', 'B', 'C', 'D', 'E'] :
##    bias_data_grids.append(np.genfromtxt(path.join(directory, 'Tmean', 'Trace21_2.5x2.5_T'+month_code+'_Delta.txt')))
##print '    Pass =', (np.array(bias_data_grids) == bias_data_grids1).all()
##
##print '  Test 2: minimum temperature, months A-L (all)'
##bias_data_grids2 = data_helper.loadBiasCorrectionDataGrids(parameter='minimum_temperature', month_indices=range(12))
##bias_data_grids = [] # expected
##for month_code in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L'] :
##    bias_data_grids.append(np.genfromtxt(path.join(directory, 'Tmin', 'Trace21_2.5x2.5_I'+month_code+'_Delta.txt')))
##print '    Pass =', (np.array(bias_data_grids) == bias_data_grids2).all()
##
##print '  Test 3: maximum temperature, months A,B,C,K,L'
##bias_data_grids3 = data_helper.loadBiasCorrectionDataGrids(parameter='maximum_temperature', month_indices=[0,1,2,10,11])
##bias_data_grids = [] # expected
##for month_code in ['K', 'L', 'A', 'B', 'C'] :
##    bias_data_grids.append(np.genfromtxt(path.join(directory, 'Tmax', 'Trace21_2.5x2.5_A'+month_code+'_Delta.txt')))
##print '    Pass =', (np.array(bias_data_grids) == bias_data_grids3).all()
##
##print '  Test 4: relative_humidity, months A,B,C,K,L'
##bias_data_grids4 = data_helper.loadBiasCorrectionDataGrids(parameter='relative_humidity', month_indices=[0,1,2,10,11])
##bias_data_grids = [] # expected
##for month_code in ['K', 'L', 'A', 'B', 'C'] :
##    bias_data_grids.append([np.genfromtxt(path.join(directory, 'RelativeHumidity', 'Trace21_2.5x2.5_H'+month_code+'_Modelled.txt')),
##                            np.genfromtxt(path.join(directory, 'RelativeHumidity', 'Trace21_2.5x2.5_H'+month_code+'_Observed.txt'))])
##print '    Pass =', (np.array(bias_data_grids) == bias_data_grids4).all()
##
##print '  Test 6: precipitation, month G'
##bias_data_grids6 = data_helper.loadBiasCorrectionDataGrids(parameter='precipitation', month_indices=[6])
##bias_data_grids = [np.genfromtxt(path.join(directory, 'Precip', 'Trace21_2.5x2.5_PG_Error.txt'))]
##print '    Pass =', (np.array(bias_data_grids) == bias_data_grids6).all()
##
### TEST applyBoundToDataGrid prior to bias correction
##print 'TEST applyBoundToDataGrid:'
##print '  Test 1: lower bound of 0.0 for test 6: precipitation'
##within_bounds_indices = (data_grids6 >= 0.0).nonzero()
##out_of_bounds_indices = (data_grids6 < 0.0).nonzero()
##bounded_data_grids6 = data_helper.applyBoundToDataGrid(data_grids6, 'lower', 0.0)
##print '    Pass =', ( bounded_data_grids6.min() >= 0.0 and
##                      (bounded_data_grids6[out_of_bounds_indices] == 0.0).all() and
##                      (bounded_data_grids6[within_bounds_indices] == data_grids6[within_bounds_indices]).all() )
##print '  Test 2: upper bound of 100.0 for test 4: relative_humidity'
##within_bounds_indices = (data_grids4 <= 100.0).nonzero()
##out_of_bounds_indices = (data_grids4 > 100.0).nonzero()
##bounded_data_grids4 = data_helper.applyBoundToDataGrid(data_grids4, 'upper', 100.0)
##print '    Pass =', ( bounded_data_grids4.max() <= 100.0 and
##                      (bounded_data_grids4[out_of_bounds_indices] == 100.0).all() and
##                      (bounded_data_grids4[within_bounds_indices] == data_grids4[within_bounds_indices]).all() )
##
### TEST loadClimateDataGrids with bias correction
##print 'TEST loadClimateDataGrids with bias correction:'
###data_helper.setClimateDataDirectory(r'C:/Users/shaythorne/Desktop/PaleoView/Tests/NetCdfClimateData')
###data_helper.use_netCdf_data = True
##directory = data_helper.climate_data_directory['path']
##
##print '  Test 1: mean temperature, 1989 AD, months A-E'
##corrected_data_grids1 = data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=1989, month_indices=range(5), correct_bias=True)
##corrected_data_grids = data_grids1 + bias_data_grids1
##print '    Pass =', (corrected_data_grids == corrected_data_grids1).all()
##
##print '  Test 2: minimum temperature, 10 BP (1940), months A-L (all)'
##corrected_data_grids2 = data_helper.loadClimateDataGrids(parameter='minimum_temperature', year_ad=1940, month_indices=range(12), correct_bias=True)
##corrected_data_grids = data_grids2 + bias_data_grids2
##print '    Pass =', (corrected_data_grids == corrected_data_grids2).all()
##
##print '  Test 3: maximum temperature, 1989 AD, months A,B,C,K,L'
##corrected_data_grids3 = data_helper.loadClimateDataGrids(parameter='maximum_temperature', year_ad=1989, month_indices=[0,1,2,10,11], correct_bias=True)
##corrected_data_grids = data_grids3 + bias_data_grids3
##print '    Pass =', (corrected_data_grids == corrected_data_grids3).all()
##
##print '  Test 4: relative_humidity, 1951 AD, months A,B,C,K,L'
##corrected_data_grids4 = data_helper.loadClimateDataGrids(parameter='relative_humidity', year_ad=1951, month_indices=[0,1,2,10,11], correct_bias=True)
##corrected_data_grids = np.zeros((5, data_helper.grid_height, data_helper.grid_width))
##for i in range(5) :
##    for row in range(data_helper.grid_height) :
##        for col in range(data_helper.grid_width) :
##            M1uncorr = bounded_data_grids4[i,row,col]
##            Mref = bias_data_grids4[i,0,row,col]
##            Oref = bias_data_grids4[i,1,row,col]
##            if M1uncorr <= Mref :
##                M1corr = M1uncorr*(Oref/Mref)
##            elif M1uncorr >= Mref :
##                M1corr = Oref + ((100-Oref)/(100-Mref))*(M1uncorr-Mref)
##            corrected_data_grids[i,row,col] = M1corr
##print '    Pass =', (corrected_data_grids == corrected_data_grids4).all()
##
##print '  Test 6: precipitation, 1989 AD, month G'
##corrected_data_grids6 = data_helper.loadClimateDataGrids(parameter='precipitation', year_ad=1989, month_indices=[6], correct_bias=True)
##corrected_data_grids = data_grids6*bias_data_grids6
##print '    Pass =', (corrected_data_grids.round(6) == corrected_data_grids6.round(6)).all() # bias_data_grids6 contains inf's
##
### TEST rearrangeMonthIndices
##print 'TEST rearrangeMonthIndices:'
##print '  Test 1: [0,1,2,10,11] => [10,11,0,1,2]'
##print '    Pass =', data_helper.rearrangeMonthIndices([0,1,2,10,11]) == [10,11,0,1,2]
##print '  Test 2: [10,11,0,1,2] => [10,11,0,1,2]'
##print '    Pass =', data_helper.rearrangeMonthIndices([10,11,0,1,2]) == [10,11,0,1,2]
##print '  Test 3: [0,10,2,11,1] => [10,11,0,1,2]'
##print '    Pass =', data_helper.rearrangeMonthIndices([0,10,2,11,1]) == [10,11,0,1,2]
##print '  Test 4: [0,1,2] => [0,1,2]'
##print '    Pass =', data_helper.rearrangeMonthIndices([0,1,2]) == [0,1,2]
##print '  Test 5: [10,11] => [10,11]'
##print '    Pass =', data_helper.rearrangeMonthIndices([10,11]) == [10,11]
##print '  Test 6: range(12) => range(12)'
##print '    Pass =', data_helper.rearrangeMonthIndices(range(12)) == range(12)
##
##
### TEST generateParameterData with bias correction
##print 'TEST generateParameterData with bias correction:'
##
##print '  Test 1: minimum temperature, 1950-1980, interval steps 10 yr, size 10 yr, months B, region full, grid, averaged'
##corrected_parameter_data1 = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                              parameter_code='minimum-temperature',
##                                                              period_ad_from=1950,
##                                                              period_ad_until=1980,
##                                                              interval_step=10,
##                                                              interval_size=10,
##                                                              month_indices=[1],
##                                                              region_mask=1,
##                                                              generate_grids=True,
##                                                              all_months=False,
##                                                              correct_bias=True)
##bias_data_for_parameter_data1 = data_helper.loadBiasCorrectionDataGrids(parameter='minimum_temperature', month_indices=[1])
##print '    Pass =', ((np.array(parameter_data1) + bias_data_for_parameter_data1[0]).round(12) == np.array(corrected_parameter_data1).round(12)).all()
##
##region_mask = np.zeros((data_helper.grid_height, data_helper.grid_width))
##region_mask[0:4,0:4] = 1
##
##print '  Test 4: minimum temperature, 1950-1980, interval steps 10 yr, size 10 yr, all months, region TL4x4, non-grid'
##corrected_parameter_data4 = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                              parameter_code='minimum-temperature',
##                                                              period_ad_from=1950,
##                                                              period_ad_until=1980,
##                                                              interval_step=10,
##                                                              interval_size=10,
##                                                              month_indices=range(12),
##                                                              region_mask=region_mask,
##                                                              generate_grids=False,
##                                                              all_months=True,
##                                                              correct_bias=True)
##bias_data_for_parameter_data4 = data_helper.loadBiasCorrectionDataGrids(parameter='minimum_temperature', month_indices=range(12))
##stats_data = []
##for month_index in range(12) :
##    data_grids_month = []
##    region_masks = []
##    for from_year_ad in range(1950, 1981, 10) :
##        interval_grids = []
##        for year_ad in range(from_year_ad-5, from_year_ad+5) :
##            grid = data_helper.loadClimateDataGrids(parameter='minimum_temperature', year_ad=year_ad, month_indices=[month_index]) + bias_data_for_parameter_data4[month_index,:,:]
##            interval_grids.append(grid)
##        data_grids_month.append(np.array(interval_grids).mean(1).mean(0))
##        region_masks.append(region_mask)
##    stats_data.append(data_helper.calculateGridRegionStatistics(data_grids_month, region_masks))
##print '    Pass =', (stats_data == corrected_parameter_data4)
##
### TEST generateParameterData with time dependent region masks
##print 'TEST generateParameterData with time dependent region masks:'
##
##print '  Test 1: mean temperature, 1760-1960, interval steps 100 yr, size 10 yr, months A, region time dependent land, non-grid'
##time_dependent_region_masks = data_helper.loadRegionMask('land-0-21KBP', time_dependent=True)
##time_dependent_regions_parameter_data1 = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                                           parameter_code='mean-temperature',
##                                                                           period_ad_from=1760,
##                                                                           period_ad_until=1960,
##                                                                           interval_step=100,
##                                                                           interval_size=10,
##                                                                           month_indices=[0],
##                                                                           region_mask=time_dependent_region_masks,
##                                                                           generate_grids=False,
##                                                                           all_months=False)
##data_grids = []
##region_masks = []
##region_mask_indices = [200, 100, 0]
##for i, from_year_ad in enumerate(range(1760, 1961, 100)) :
##    interval_grids = []
##    for year_ad in range(from_year_ad-5, from_year_ad+5) :
##        interval_grids.append(data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=year_ad, month_indices=[0]))
##    data_grids.append(np.array(interval_grids).mean(1).mean(0))
##    region_masks.append(time_dependent_region_masks[region_mask_indices[i]])
##stats_data = data_helper.calculateGridRegionStatistics(data_grids, region_masks)
##print '    Pass =', (stats_data == time_dependent_regions_parameter_data1)
##
##print '  Test 2: mean temperature, 1760-1960, interval steps 100 yr, size 10 yr, all months, region time dependent land, non-grid'
##time_dependent_regions_parameter_data2 = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                                           parameter_code='mean-temperature',
##                                                                           period_ad_from=1760,
##                                                                           period_ad_until=1960,
##                                                                           interval_step=100,
##                                                                           interval_size=10,
##                                                                           month_indices=range(12),
##                                                                           region_mask=time_dependent_region_masks,
##                                                                           generate_grids=False,
##                                                                           all_months=True)
##stats_data = []
##for month_index in range(12) :
##    data_grids_month = []
##    region_masks = []
##    region_mask_indices = [200, 100, 0]
##    for i, from_year_ad in enumerate(range(1760, 1961, 100)) :
##        interval_grids = []
##        for year_ad in range(from_year_ad-5, from_year_ad+5) :
##            interval_grids.append(data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=year_ad, month_indices=[month_index]))
##        data_grids_month.append(np.array(interval_grids).mean(1).mean(0))
##        region_masks.append(time_dependent_region_masks[region_mask_indices[i]])
##    stats_data.append(data_helper.calculateGridRegionStatistics(data_grids_month, region_masks))
##print '    Pass =', (stats_data == time_dependent_regions_parameter_data2)
##
### Climate Data via URL
##data_helper.setClimateDataSource('url')
##data_helper.setClimateDataUrl('http://homepage.cs.latrobe.edu.au/shaythorne/paleoview/')
##
### TEST loadClimateDataGrids via URL
##print 'TEST loadClimateDataGrids via URL:'
##directory = data_helper.climate_data_directory['path']
##
##print '  Test 1: mean temperature, 1989 AD, months A-E'
##data_grids1 = data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=1989, month_indices=range(5))
##data_grids = [] # expected
##for month_code in ['A', 'B', 'C', 'D', 'E'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'T', 'Trace21_2.5x2.5_1989AD.1T'+month_code+'.txt')))  
##print '    Pass =', (np.array(data_grids) == data_grids1).all()
##
##print '  Test 2: minimum temperature, 10 BP (1940), months A-L (all)'
##data_grids2 = data_helper.loadClimateDataGrids(parameter='minimum_temperature', year_ad=1940, month_indices=range(12))
##data_grids = [] # expected
##for month_code in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'Tmin', 'Trace21_2.5x2.5_10BP.1I'+month_code+'.txt')))  
##print '    Pass =', (np.array(data_grids) == data_grids2).all()
##
##print '  Test 3: maximum temperature, 1989 AD, months A,B,C,K,L'
##data_grids3 = data_helper.loadClimateDataGrids(parameter='maximum_temperature', year_ad=1989, month_indices=[0,1,2,10,11])
##data_grids = [] # expected
##for month_code in ['K', 'L'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'Tmax', 'Trace21_2.5x2.5_1988AD.1A'+month_code+'.txt')))  
##for month_code in ['A', 'B', 'C'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'Tmax', 'Trace21_2.5x2.5_1989AD.1A'+month_code+'.txt')))  
##print '    Pass =', (np.array(data_grids) == data_grids3).all()
##
##print '  Test 4: relative_humidity, 1951 AD, months A,B,C,K,L'
##data_grids4 = data_helper.loadClimateDataGrids(parameter='relative_humidity', year_ad=1951, month_indices=[0,1,2,10,11])
##data_grids = [] # expected
##for month_code in ['K', 'L'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'H', 'Trace21_2.5x2.5_0BP.1H'+month_code+'.txt')))  
##for month_code in ['A', 'B', 'C'] :
##    data_grids.append(np.genfromtxt(path.join(directory, 'H', 'Trace21_2.5x2.5_1951AD.1H'+month_code+'.txt')))  
##print '    Pass =', (np.array(data_grids) == data_grids4).all()
##
##print '  Test 5: specific humidity, 1989 AD, month G'
##data_grids5 = data_helper.loadClimateDataGrids(parameter='specific_humidity', year_ad=1989, month_indices=[6])
##data_grids = [np.genfromtxt(path.join(directory, 'Q', 'Trace21_2.5x2.5_1989AD.1SG.txt'))] # expected
##print '    Pass =', (np.array(data_grids) == data_grids5).all()
##
##print '  Test 6: precipitation, 1989 AD, month G'
##data_grids6 = data_helper.loadClimateDataGrids(parameter='precipitation', year_ad=1989, month_indices=[6])
##data_grids = [np.genfromtxt(path.join(directory, 'P', 'Trace21_2.5x2.5_1989AD.1PG.txt'))] # expected
##print '    Pass =', (np.array(data_grids) == data_grids6).all()
##
##print '  Test 7: missing mean_temperature data at Dropbox URL'
##try :
##    data_grids7 = data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=100, month_indices=[6])
##except Exception, e :
##    print '    Exception:', e
##
##print '  Test 8: missing sea_level_pressure data at Dropbox URL'
##try :
##    data_grids8 = data_helper.loadClimateDataGrids(parameter='sea_level_pressure', year_ad=1989, month_indices=[6])
##except Exception, e :
##    print '    Exception:', e
##
##print '  Test 9: faulty URL '
##data_helper.setClimateDataUrl('http://homepage.cs.latrobe.edu.au/shaythorne/paleoviewBAD/')
##try :
##    data_grids9 = data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=1989, month_indices=[6])
##except Exception, e :
##    print '    Exception:', e
##data_helper.setClimateDataUrl('http://homepage.cs.latrobe.edu.au/shaythorne/paleoview/')
##
### Climate Data via URL
##data_helper.setClimateDataSource('url')
##data_helper.setClimateDataUrl('http://homepage.cs.latrobe.edu.au/shaythorne/paleoview/')
##
### TEST loadNongriddedDataFrame via URL
##print 'TEST loadNongriddedDataFrame via URL:'
##full_soi_dataframe = pd.read_csv(path.join(directory, 'South_Oscillation_Index.txt'), delim_whitespace=True)
##
##print '  Test 1: soi, 10-5 BP (1939-45)'
##data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation',
##                                    parameter_code='soi',
##                                    period_ad_from=1940,
##                                    period_ad_until=1945)
##data_frame = full_soi_dataframe[((full_soi_dataframe['BP/AD'] == 'BP') & (full_soi_dataframe['Yrs'] >= 5) & (full_soi_dataframe['Yrs'] <= 11))]
##print '    Pass =', (data_frame.get_values() == data_helper.nongridded_data_frame.get_values()).all()
##
##print '  Test 2: soi, 5 BP - 1955 AD (1944-1955)'
##data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation',
##                                    parameter_code='soi',
##                                    period_ad_from=1945,
##                                    period_ad_until=1955)
##data_frame = full_soi_dataframe[( ((full_soi_dataframe['BP/AD'] == 'BP') & (full_soi_dataframe['Yrs'] >= 0) & (full_soi_dataframe['Yrs'] <= 6)) |
##                                  ((full_soi_dataframe['BP/AD'] == 'AD') & (full_soi_dataframe['Yrs'] >= 1951) & (full_soi_dataframe['Yrs'] <= 1955)) )]
##print '    Pass =', (data_frame.get_values() == data_helper.nongridded_data_frame.get_values()).all()
##
##print '  Test 3: soi, 1979-89 AD'
##data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation',
##                                    parameter_code='soi',
##                                    period_ad_from=1980,
##                                    period_ad_until=1989)
##data_frame = full_soi_dataframe[((full_soi_dataframe['BP/AD'] == 'AD') & (full_soi_dataframe['Yrs'] >= 1979) & (full_soi_dataframe['Yrs'] <= 1989))]
##print '    Pass =', (data_frame.get_values() == data_helper.nongridded_data_frame.get_values()).all()
##
##print '  Test 4: soi, 5 BP - 1955 AD (1944-1955) and delta period 100-95BP (1849-1855)'
##data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation',
##                                    parameter_code='soi',
##                                    period_ad_from=1945,
##                                    period_ad_until=1955,
##                                    delta_interval={ 'ad_from' : 1850, 'ad_until' : 1855 })
##data_frame = full_soi_dataframe[( ((full_soi_dataframe['BP/AD'] == 'BP') & (full_soi_dataframe['Yrs'] >= 0) & (full_soi_dataframe['Yrs'] <= 6)) |
##                                  ((full_soi_dataframe['BP/AD'] == 'AD') & (full_soi_dataframe['Yrs'] >= 1951) & (full_soi_dataframe['Yrs'] <= 1955)) |
##                                  ((full_soi_dataframe['BP/AD'] == 'BP') & (full_soi_dataframe['Yrs'] >= 95) & (full_soi_dataframe['Yrs'] <= 101)) )]
##print '    Pass =', (data_frame.get_values() == data_helper.nongridded_data_frame.get_values()).all()
##
##print '  Test 5: faulty URL '
##data_helper.setClimateDataUrl('http://homepage.cs.latrobe.edu.au/shaythorne/paleoviewBAD/')
##try :
##    data_helper.loadNongriddedDataFrame(parameter_group_code='southern-oscillation',
##                                        parameter_code='soi',
##                                        period_ad_from=1945,
##                                        period_ad_until=1955)
##except Exception, e :
##    print '    Exception:', e
##data_helper.setClimateDataUrl('http://homepage.cs.latrobe.edu.au/shaythorne/paleoview/')
##
### TEST generateGridDataFile
##print 'TEST generateGridDataFile: with Northern Hemisphere mask and some nan/-inf/+inf'
##data_helper.setClimateDataSource('local')
##data_file_test_dir = data_helper.getFileGenerationDirectoryPath()
##region_mask = np.zeros((data_helper.grid_height, data_helper.grid_width))
##region_mask[0:36,:] = 1
##data_grid = data_helper.loadClimateDataGrids(parameter='mean_temperature', year_ad=1989, month_indices=[0])[0]
##data_grid[0,0] = data_grid[0,0]*0./0.
##data_grid[1,0] = data_grid[1,0]*1./0.
##data_grid[0,1] = data_grid[0,1]*-1./0.
##masked_data = np.ma.masked_array(data_grid, mask=((region_mask - 1)*-1))
##expected_data_grid = data_grid.copy()
##expected_data_grid[((region_mask - 1)*-1).nonzero()] = -9999.
##expected_data_grid[((np.isfinite(expected_data_grid)-1)*-1).nonzero()] = -9999.
##
##print '  Test 1: generate CSV'
##data_helper.generateGridDataFile(masked_data.copy(), file_type='csv', year_label='1989')
##grid_from_file = pd.read_csv(path.join(data_file_test_dir, 'grid_data_1989.csv'), header=None).as_matrix()
##print '    Pass =', (grid_from_file.round(5) == expected_data_grid.round(5)).all()
##
##print '  Test 2: generate ASCII'
##data_helper.generateGridDataFile(masked_data.copy(), file_type='ascii', year_label='1989')
##grid_from_file = pd.read_fwf(path.join(data_file_test_dir, 'grid_data_1989.txt'), header=None).as_matrix()
##print '    Pass =', (grid_from_file.round(5) == expected_data_grid.round(5)).all()
##
##print '  Test 3: generate ESRI ASCII'
##data_helper.generateGridDataFile(masked_data.copy(), file_type='esri_ascii', year_label='1989')
##f = open(path.join(data_file_test_dir, 'grid_data_1989.asc'))
##header_from_file = f.readlines()[:6]
##f.close()
##grid_from_file = pd.read_table(path.join(data_file_test_dir, 'grid_data_1989.asc'), sep=' ', skiprows=6, header=None).as_matrix()
##print '    Pass =', (grid_from_file.round(5) == expected_data_grid.round(5)).all() and header_from_file == ['ncols 144\n', 'nrows 72\n', 'xllcorner -180\n', 'yllcorner -90\n', 'cellsize 2.5\n', 'nodata_value -9999\n']
##
##expected_masked_data = masked_data.copy()
##expected_masked_data[((np.isfinite(expected_masked_data)-1)*-1).nonzero()] = np.nan
##
##print '  Test 5: generate netCDF'
##data_helper.generateGridDataFile(masked_data.copy(), file_type='netcdf', year_label='1989', description='test descr', data_units='test units')
##rootgrp = Dataset(path.join(data_file_test_dir, 'grid_data_1989.nc'))
###grid_from_file = pd.read_fwf(path.join(data_file_test_dir, 'grid_data_1989.nc'), header=None).as_matrix()
##print '    Pass =', (np.isnan(rootgrp.variables['data'][0,0:2]).all() and (rootgrp.variables['data'][0,2:] == expected_masked_data[0,2:]).all() and
##                     np.isnan(rootgrp.variables['data'][1,0:1]).all() and (rootgrp.variables['data'][1,1:] == expected_masked_data[1,1:]).all() and
##                     (rootgrp.variables['data'][2:,:] == expected_masked_data[2:,:]).all() and
##                     rootgrp.description == 'test descr' and rootgrp.variables['data'].units == 'test units' and
##                     (rootgrp.variables['latitudes'][:] == np.arange(88.75,-88.751,-2.5)).all() and rootgrp.variables['latitudes'].units == 'degrees north' and
##                     (rootgrp.variables['longitudes'][:] == np.arange(-178.75,178.751,2.5)).all() and rootgrp.variables['longitudes'].units == 'degrees east')
##
### TEST generateSeriesDataFile
##print 'TEST generateSeriesDataFile:'
##data_helper.setClimateDataSource('local')
##data_file_test_dir = data_helper.getFileGenerationDirectoryPath()
##statistics_data = data_helper.generateParameterData(parameter_group_code='temperature',
##                                                    parameter_code='mean-temperature',
##                                                    period_ad_from=1930,
##                                                    period_ad_until=1980,
##                                                    interval_step=10,
##                                                    interval_size=10,
##                                                    month_indices=[0],
##                                                    region_mask=1,
##                                                    generate_grids=False,
##                                                    all_months=False)
##statistics_data['year'] = [1930, 1940, 1950, 1960, 1970, 1980]
##stats_data_frame = pd.DataFrame(statistics_data)[['year', 'minimum', 'area_mean', 'maximum']]
##
##print '  Test 1: generate CSV'
##data_helper.generateSeriesDataFile(data_frame=stats_data_frame, file_type='csv')
##frame_from_file = pd.read_csv(path.join(data_file_test_dir, 'series_data.csv'))
##print '    Pass =', frame_from_file.to_string() == stats_data_frame.to_string()
##
##print '  Test 2: generate ASCII'
##data_helper.generateSeriesDataFile(data_frame=stats_data_frame, file_type='ascii', month='jan')
##frame_from_file = pd.read_fwf(path.join(data_file_test_dir, 'series_data_jan.txt'))
##print '    Pass =', frame_from_file.to_string() == stats_data_frame.to_string()
##
##print '  Test 3: generate netCDF'
##data_helper.generateSeriesDataFile(data_frame=stats_data_frame, file_type='netcdf', description='test descr', data_units='test units')
##rootgrp = Dataset(path.join(data_file_test_dir, 'series_data.nc'))
##print '    Pass =', ((rootgrp.variables['year'][:] == np.array(statistics_data['year'])).all() and rootgrp.description == 'test descr' and
##                     (rootgrp.variables['minimum'][:] == np.array(statistics_data['minimum'])).all() and rootgrp.variables['minimum'].units == 'test units' and
##                     (rootgrp.variables['area_mean'][:] == np.array(statistics_data['area_mean'])).all() and rootgrp.variables['area_mean'].units == 'test units' and
##                     (rootgrp.variables['maximum'][:] == np.array(statistics_data['maximum'])).all() and rootgrp.variables['maximum'].units == 'test units')
##
### TEST generateWordTable
##print 'TEST generateWordTable:'
##data_frame = pd.DataFrame({ 'Column 1' : np.arange(3), 'Column 2' : np.arange(3) + 2.0**0.5, 'Column 3' : np.arange(3)*100.0 })
##file_path = path.join(data_helper.getFileGenerationDirectoryPath() ,'test_table.docx')
##data_helper.generateWordTable(data_frame, 'Title Test', file_path=file_path)
##document = Document(file_path)
##print '  Pass =', ( document.paragraphs[0].text == 'Title Test' and
##                    document.tables[0].cell(0,0).text == 'Column 1' and document.tables[0].cell(0,1).text == 'Column 2' and document.tables[0].cell(0,2).text == 'Column 3' and
##                    document.tables[0].cell(1,0).text == '0' and document.tables[0].cell(1,1).text == str(2.0**0.5) and document.tables[0].cell(1,2).text == '0.0' and
##                    document.tables[0].cell(2,0).text == '1' and document.tables[0].cell(2,1).text == str(1+2.0**0.5) and document.tables[0].cell(2,2).text == '100.0' and
##                    document.tables[0].cell(3,0).text == '2' and document.tables[0].cell(3,1).text == str(2+2.0**0.5) and document.tables[0].cell(3,2).text == '200.0' )
##
### TEST generateTextTable
##print 'TEST generateTextTable:'
##data_frame = pd.DataFrame({ 'Column 1' : np.arange(3), 'Column 2' : np.arange(3) + 2.0**0.5, 'Column 3' : np.arange(3)*100.0 })
##file_path = path.join(data_helper.getFileGenerationDirectoryPath() ,'test_table.txt')
##data_helper.generateTextTable(data_frame, 'Title Test', file_path=file_path)
##f = open(file_path)
##title_line = f.readline()
##title_underline = f.readline()
##f.close()
##frame_from_file = pd.read_fwf(file_path, skiprows=2)
##print '    Pass =', title_line == ('Title Test\n') and title_underline == ('----------\n') and frame_from_file.to_string(float_format=(lambda f: '%.3f'%f)) == data_frame.to_string(float_format=(lambda f: '%.3f'%f))
##
##print STOP
##
### TEST generateNetCdfClimateData
##print 'TEST generateNetCdfClimateData:'
### Set the climate data and output directories
##data_helper.setClimateDataSource('local')
##data_helper.setClimateDataDirectory(r'D:/Trace-21') # 
##data_helper.useNetCdfData(False)
##data_helper.setFileGenerationDirectory(r'C:/Users/shaythorne/Desktop/PaleoView/Tests/DownloadedClimateData') # NetCdfClimateData
####parameter = 'mean_temperature'
####from_year_ad = 1950 - 22000
####until_year_ad = 1950 - 15000
####last_grids = data_helper.generateNetCdfClimateData(parameter, from_year_ad, until_year_ad, zlib=True)
##parameters = ['relative_humidity'] #'mean_temperature', 'minimum_temperature', 'maximum_temperature', 'specific_humidity', 'precipitation', 'sea_level_pressure']
###intervals = [[1950-22000, 1950-15000], [1950-15000, 1950-10000], [1950-10000, 1950-5000], [1950-5000, 1989]] # 
##intervals = [[1700, 1989]] # 
##for parameter in parameters :
##    for interval in intervals :
##        data_helper.generateNetCdfClimateData(parameter, interval[0], interval[1], zlib=True, decimals=6) #, correction_factor=0.001**bad decimals)
##        for year_ad in range(interval[0], interval[1]+1) :
##            for month_index in range(12) :
##                data_file_present = data_helper.climateDataFilePresent(parameter, year_ad, month_index)
##                if not data_file_present['present'] :
##                    print data_file_present['data_file'], 'is missing'
##
##print STOP
##
### TEST downloadClimateData
##print 'TEST downloadClimateData:'
##data_helper.setClimateDataSource('local')
##data_helper.setClimateDataDirectory(r'C:/Users/shaythorne/Desktop/PaleoView/Tests/DownloadedClimateData')
##data_helper.setClimateDataUrl('http://homepage.cs.latrobe.edu.au/shaythorne/paleoview/')
##parameter = 'mean_temperature'
##from_year_ad = 1950 - 5000
##until_year_ad = 1989
##data_helper.downloadClimateData(parameter, from_year_ad, until_year_ad, retain_netCdf_file=True, unpack=True)

### relative humidity / 1000
##print 'relative humidity / 1000'
##data_helper.setClimateDataSource('local')
##data_helper.useNetCdfData(False)
##data_helper.setClimateDataDirectory(r'C:/Users/shaythorne/Desktop/PaleoView/Tests/NetCdfClimateData')
##data_helper.setFileGenerationDirectory(r'C:/Users/shaythorne/Desktop/PaleoView/Tests/NetCdfClimateData')
##for year_ad in range((1950-22000), 1990) :
##    for month_index in range(12) :
##        data_grid = data_helper.loadClimateDataGrid('relative_humidity', year_ad, month_index)
##        postfix = 'AD'
##        year = year_ad
##        if year_ad <= 1950 :
##            postfix = 'BP'
##            year = 1950 - year_ad
##        year_str = str(year) + postfix
##        data_file = data_helper.climate_data_file_template.replace('{parameter_directory_code}', 'Hdiv1000')
##        data_file = data_file.replace('{year}', str(year))
##        data_file = data_file.replace('{postfix}', postfix)
##        data_file = data_file.replace('{parameter_file_code}', data_helper.parameter_file_code_map['relative_humidity'])
##        data_file = data_file.replace('{month_code}', data_helper.month_codes[month_index])
##        data_file = path.join(data_helper.climate_data_directory['path'], data_file)
##        np.savetxt(data_file, data_grid/1000, fmt='%.6f', delimiter=' ')
##    if not year % 1000  :
##        print str(year) + postfix
        


