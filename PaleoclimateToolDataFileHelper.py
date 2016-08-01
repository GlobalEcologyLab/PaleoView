# Python modules
import re
import string
import urllib
import urllib2 as url
from math import floor
from os import listdir, mkdir, path, remove
from StringIO import StringIO
from time import time, localtime, strftime

# Python extension modules (requires extension installation)
import numpy as np
import pandas as pd
import netCDF4_utils
import netcdftime
from netCDF4 import Dataset
import docx
from docx import Document
from docx.enum.section import WD_ORIENT
import lxml.etree
import lxml._elementpath

## Paleoclimate Tool Data File Helper
## * Loads climate data files into arrays for the PaleoView tool.
## * Loads bias correction data for the tool.
## * Calculates delta values, that is the change in climate data values relative to a specified year
## * Calculates climate data statistics for a specified grid region
## * Generates grid data files based on aggregated climate data
## * Generates series data files based on aggregated climate data
## * Generates a text table for recording statistics
## * Generates NetCDF files for subsets of the climate data (for web-based repository)
## * Downloads and unpacks climate data subsets from the web-based NetCDF repository 
class PaleoclimateToolDataFileHelper :

    # Initialise
    def __init__(self, application_gui=None) :

        # Set the application GUI
        self.application_gui = application_gui

        # Climate data source
        self.climate_data_source = 'local' # url or local

        # Climate data url and optional proxy details
        self.climate_data_url = ''
        self.climate_data_proxy = { 'active' : False, 'url' : '', 'username' : '', 'password' : '' }

        # Climate data directory
        self.climate_data_directory = { 'name' : '', 'directory' : '', 'path' : '' }

        # Region mask directory
        self.region_mask_directory = { 'name' : '', 'directory' : '', 'path' : '' }

        # Bias correction directory
        self.bias_correction_directory = { 'name' : '', 'directory' : '', 'path' : '' }

        # File Generation directory
        self.file_generation_directory = { 'name' : '', 'directory' : '', 'path' : '' }

        # Month codes and names
        self.month_codes = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L']
        self.month_names = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']

        # Grid size
        self.grid_height = 72
        self.grid_width = 144

        # Climate data parameters in presented order
        self.data_parameters = ['mean_temperature', 'minimum_temperature', 'maximum_temperature',
                                'specific_humidity', 'relative_humidity',
                                'precipitation', 'sea_level_pressure']

        # Parameter directory code map
        self.parameter_directory_code_map = { 'mean_temperature' : 'T', 'minimum_temperature' : 'Tmin', 'maximum_temperature' : 'Tmax',
                                              'specific_humidity' : 'Q', 'relative_humidity' : 'H',
                                              'precipitation' : 'P', 'sea_level_pressure' : 'PSL' }
        # Parameter file code map
        self.parameter_file_code_map = { 'mean_temperature' : 'T', 'minimum_temperature' : 'I', 'maximum_temperature' : 'A',
                                         'specific_humidity' : 'S', 'relative_humidity' : 'H',
                                         'precipitation' : 'P', 'sea_level_pressure' : 'M' }

        # Parameter unit strings
        self.parameter_unit_string = { 'mean_temperature' : 'degrees C', 'minimum_temperature' : 'degrees C', 'maximum_temperature' : 'degrees C',
                                       'specific_humidity' : 'gm/kg', 'relative_humidity' : '%',
                                       'precipitation' : 'mm/day', 'sea_level_pressure' : 'hPa' }

        # Climate data file template
        self.climate_data_file_template = path.join('{parameter_directory_code}', 'Trace21_2.5x2.5_{year}{postfix}.1{parameter_file_code}{month_code}.txt')

        # Climate data download intervals
        self.climate_data_download_intervals = ['22000BP-15000BP', '15000BP-10000BP', '10000BP-5000BP', '5000BP-1989AD']
        self.download_data_window = 100 # needs to match with the tool's maximum interval size

        # Use downloaded NetCDF files. Maintain a current list of available files for each parameter
        self.use_netCdf_data = True
        self.current_netCdf_data_intervals = self.getCurrentNetCdfDataIntervals()

        # Cached NetCDF file and subgroup for parameter # TODO clear NetCDF data cache
        self.cached_netCdf_data = {} # { parameter : { 'rootgrp' : <DataSet>, 'root_interval_str' : str, 'sub_interval_str' : str  } }

        # Non-gridded parameter files
        self.non_gridded_parameter_files = { 'southern-oscillation' : { 'soi' : 'South_Oscillation_Index.txt', 'enso' : '' } }
        self.non_gridded_parameter_file_columns = { 'southern-oscillation' : { 'soi' : { 'postfix' : 'BP/AD', 'year' : 'Yrs', 'month_code' : 'Months', 'data' : 'SOI'},
                                                                               'enso' : { 'postfix' : 'BP/AD', 'year' : 'Yrs', 'month_code' : 'Months', 'data' : 'ENSO'} } }

        # Non-gridded parameter panda data frame (avoids loading same file multiple times)
        self.nongridded_data_frame = None

        # Parameter calculations
        self.parameter_calculation = { 'temperature' : {}, 'precipitation' : {}, 'humidity' : {}, 'sea-level-pressure' : {}, 'southern-oscillation' : {} }
        self.parameter_calculation['temperature']['mean-temperature'] = { 'parameters' : ['mean_temperature'], 'calculate' : { 'average' : 'mean_temperature' }, 'return' : 'average' }
        self.parameter_calculation['temperature']['minimum-temperature'] = { 'parameters' : ['minimum_temperature'], 'calculate' : { 'average' : 'minimum_temperature' }, 'return' : 'average' }
        self.parameter_calculation['temperature']['maximum-temperature'] = { 'parameters' : ['maximum_temperature'], 'calculate' : { 'average' : 'maximum_temperature' }, 'return' : 'average' }
        self.parameter_calculation['temperature']['diurnal-temperature-range'] = { 'parameters' : ['minimum_temperature', 'maximum_temperature'], 'calculate' : { 'average' : 'maximum_temperature - minimum_temperature' }, 'return' : 'average' }
        self.parameter_calculation['temperature']['annual-temperature-range'] = { 'parameters' : ['mean_temperature'], 'calculate' : { 'annual_range' : 'mean_temperature' }, 'return' : 'annual_range' }
        self.parameter_calculation['temperature']['isothermality'] = { 'parameters' : ['minimum_temperature', 'maximum_temperature', 'mean_temperature'], 'calculate' : { 'average' : 'maximum_temperature - minimum_temperature', 'annual_range' : 'mean_temperature' }, 'return' : 'average/annual_range*100' }
        self.parameter_calculation['temperature']['temperature-seasonality'] = { 'parameters' : ['mean_temperature'], 'calculate' : { 'stdev_seasonality' : 'mean_temperature' }, 'return' : 'stdev_seasonality*100' }
        self.parameter_calculation['precipitation']['mean-precipitation'] = { 'parameters' : ['precipitation'], 'calculate' : { 'average' : 'precipitation' }, 'return' : 'average' }
        self.parameter_calculation['precipitation']['precipitation-seasonality'] = { 'parameters' : ['precipitation'], 'calculate' : { 'coeff_var_seasonality' : 'precipitation' }, 'return' : 'coeff_var_seasonality' }
        self.parameter_calculation['humidity']['specific-humidity'] = { 'parameters' : ['specific_humidity'], 'calculate' : { 'average' : 'specific_humidity' }, 'return' : 'average' }
        self.parameter_calculation['humidity']['relative-humidity'] = { 'parameters' : ['relative_humidity'], 'calculate' : { 'average' : 'relative_humidity' }, 'return' : 'average' }
        self.parameter_calculation['sea-level-pressure']['sea-level-pressure'] = { 'parameters' : ['sea_level_pressure'], 'calculate' : { 'average' : 'sea_level_pressure' }, 'return' : 'average' }
        self.parameter_calculation['southern-oscillation']['soi'] = { 'parameters' : ['soi'], 'calculate' : { 'average' : 'soi' }, 'return' : 'average' }
        self.parameter_calculation['southern-oscillation']['enso'] = { 'parameters' : ['enso'], 'calculate' : { 'average' : 'enso' }, 'return' : 'average' }

        # Cached bias correction data grids (avoids loading same files multiple times)
        self.cached_bias_correction_data_grids = {} # for each parameter and month

        # Bias correction directory code map
        self.bias_correction_directory_code_map = { 'mean_temperature' : 'Tmean', 'minimum_temperature' : 'Tmin', 'maximum_temperature' : 'Tmax',
                                                    'specific_humidity' : 'SpecificHumidity', 'relative_humidity' : 'RelativeHumidity',
                                                    'precipitation' : 'Precip', 'sea_level_pressure' : 'MSLP' }
        # Bias correction file code map
        self.bias_correction_file_code_map = { 'mean_temperature' : 'Delta', 'minimum_temperature' : 'Delta', 'maximum_temperature' : 'Delta',
                                               'specific_humidity' : 'Delta', 'relative_humidity' : 'Delta',
                                               'precipitation' : 'Error', 'sea_level_pressure' : 'Delta' }

        # Bias correction file template
        self.bias_correction_file_template = path.join('{bias_correction_directory_code}', 'Trace21_2.5x2.5_{parameter_file_code}{month_code}_{bias_correction_file_code}.txt')

        # pre-bias correction bounds
        self.pre_bias_correction_bounds = { 'relative_humidity' : { 'climate_data_grids' : { 'upper' : 100.0 } } }

        # Bias correction calculations
        self.bias_correction_calculation = { 'mean_temperature' : 'climate_data_grids + bias_correction_data_grids',
                                             'minimum_temperature' : 'climate_data_grids + bias_correction_data_grids',
                                             'maximum_temperature' : 'climate_data_grids + bias_correction_data_grids',
                                             'specific_humidity' : 'climate_data_grids*bias_correction_data_grids',
                                             'relative_humidity' : '(1 - (1 - climate_data_grids/100)**bias_correction_data_grids)*100',
                                             'precipitation' : 'climate_data_grids*bias_correction_data_grids',
                                             'sea_level_pressure' : 'climate_data_grids + bias_correction_data_grids' }

    # Get data parameters
    def getDataParameters(self) :
        return self.data_parameters[:]

    # Get data parameters required
    def getDataParametersRequired(self, parameter_group, parameter_name) :
        return self.parameter_calculation[parameter_group][parameter_name]['parameters']
    
    # Get climate data download intervals
    def getClimateDataDownloadIntervals(self) :
        return self.climate_data_download_intervals[:]

    # Use downloaded NetCDF data files
    def useNetCdfData(self, use=None) :
        if use != None :
            self.use_netCdf_data = use
        return self.use_netCdf_data

    # Get the month codes
    def getMonthCodes(self) :
        return self.month_codes[:]

    # Set climate data source
    def setClimateDataSource(self, code) : # url or local
        self.climate_data_source = code

    # Set climate data url
    def setClimateDataUrl(self, url) :
        if url and url[-1] != '/' :
            url += '/'
        self.climate_data_url = url

    # Set climate data proxy
    def setClimateDataProxy(self, active, url=None, username=None, password=None) :
        self.climate_data_proxy['active'] = active
        if active :
            if url != None :
                self.climate_data_proxy['url'] = url
            if username != None :
                self.climate_data_proxy['username'] = username
            if password != None :
                self.climate_data_proxy['password'] = password

    # Setup proxy for web-based climate data when required
    def setupClimateDataProxy(self) :
        if self.climate_data_proxy['active'] and self.climate_data_proxy['url'].replace('http://','').find(':') :
            proxy_host, proxy_port = self.climate_data_proxy['url'].replace('http://','').replace('/','').split(':')
            if self.climate_data_proxy['username'] and self.climate_data_proxy['password'] :
                proxy_handler = url.ProxyHandler({ 'http' : 'http://'+self.climate_data_proxy['username']+':'+self.climate_data_proxy['password']+'@'+proxy_host+':'+proxy_port })
            else :
                proxy_handler = url.ProxyHandler({ 'http' : 'http://'+proxy_host+':'+proxy_port })
            proxy_opener = url.build_opener(proxy_handler)
            url.install_opener(proxy_opener)

    # Set climate data directory
    def setClimateDataDirectory(self, path) :
        self.climate_data_directory = self.splitPath(path)

    # Get climate data directory path
    def getClimateDataDirectoryPath(self) :
        return self.climate_data_directory['path']

    # Get climate data url
    def getClimateDataUrl(self) :
        return self.climate_data_url

    # Check climate data url
    def checkClimateDataUrl(self) :
        try :
            current_url_file = url.urlopen((self.climate_data_url + 'current_url.txt'))
            current_url_file.close()
            return True
        except Exception, e :
            current_url_file = url.urlopen('http://homepage.cs.latrobe.edu.au/shaythorne/paleoview/current_url.txt')
            self.climate_data_url = current_url_file.readline()
            current_url_file.close()
            return False

    # Climate data is present
    def climateDataIsPresent(self, parameters=['any'], years={ 'from_year_ad' : None, 'until_year_ad' : None }) :
        if self.climate_data_source == 'local' :
            if parameters == ['any'] :
                climate_data_found = False
                for parameter, directory_code in self.parameter_directory_code_map.items() :
                    if self.use_netCdf_data :
                        for interval_label in self.climateDataDownloadIntervalsRequired({ 'from_year_ad' : (1950-22000), 'until_year_ad' : 1989 }) :
                            climate_data_found = (climate_data_found or self.climateDataDownloadIntervalPresent(parameter, interval_label))
                    else : # raw data
                        climate_data_found = (climate_data_found or path.exists(path.join(self.climate_data_directory['path'], directory_code)))
            elif parameters :
                climate_data_found = []
                if years['from_year_ad'] != None and years['until_year_ad'] != None :
                    for parameter in parameters :
                        if self.use_netCdf_data :
                            for interval_label in self.climateDataDownloadIntervalsRequired(years) :
                                climate_data_found.append(self.climateDataDownloadIntervalPresent(parameter, interval_label))
                        else : # raw data only checks directory presence for parameters
                            climate_data_found.append(path.exists(path.join(self.climate_data_directory['path'], self.parameter_directory_code_map[parameter])))
                    climate_data_found = np.array(climate_data_found).all()
                else :
                    for parameter in parameters :
                        if self.use_netCdf_data :
                            parameter_data_found = False
                            if self.use_netCdf_data : # Check for NetCDF data file
                                for interval_label in self.climateDataDownloadIntervalsRequired({ 'from_year_ad' : (1950-22000), 'until_year_ad' : 1989 }) :
                                    parameter_data_found = (parameter_data_found or self.climateDataDownloadIntervalPresent(parameter, interval_label))
                            climate_data_found.append(parameter_data_found)
                        else : # raw data only checks directory presence for parameters
                            climate_data_found.append(path.exists(path.join(self.climate_data_directory['path'], self.parameter_directory_code_map[parameter])))
                climate_data_found = np.array(climate_data_found).all()
            return climate_data_found
        else : # url
            return True

    # Climate data download intervals required
    def climateDataDownloadIntervalsRequired(self, years={ 'from_year_ad' : None, 'until_year_ad' : None }) :
        download_intervals_required = []
        for interval_label in self.climate_data_download_intervals :
            download_interval = self.convertDataDownloadIntervalLabelToAD(interval_label)
            if ( (download_interval['from_year_ad'] <= years['from_year_ad'] <= download_interval['until_year_ad']) or
                 (download_interval['from_year_ad'] <= years['until_year_ad'] <= download_interval['until_year_ad']) or
                 (years['from_year_ad'] <= download_interval['from_year_ad'] <= years['until_year_ad']) ) :
                download_intervals_required.append(interval_label)
        return download_intervals_required

    # Method clears NetCDF data file cache
    def clearNetCdfDataCache(self) :
        for parameter, cached_data in self.cached_netCdf_data.items() :
            cached_data['rootgrp'].close()
            self.cached_netCdf_data.pop(parameter)

    # Climate data download interval present
    def climateDataDownloadIntervalPresent(self, parameter, interval_label) :

        if self.use_netCdf_data : # Check for NetCDF data file

            expected_netCdf_file = (parameter+'-'+interval_label+'.nc')
            return path.exists(path.join(self.climate_data_directory['path'], expected_netCdf_file))

        else : # Check for raw climate data

            # Calculate middle year
            download_interval = self.convertDataDownloadIntervalLabelToAD(interval_label)
            middle_year_ad = int((download_interval['from_year_ad'] + download_interval['until_year_ad'])/2)

            # Check for middle year (first month) data grid file
            present = self.climateDataFilePresent(parameter, middle_year_ad, 0)['present']
            return present

    # Climate data file present
    def climateDataFilePresent(self, parameter, year_ad, month_index) :
        if year_ad > 1950 :
            year_str = str(year_ad) + 'AD'
        else :
            year_str = str(1950 - year_ad) + 'BP'
        data_grid_file = self.climate_data_file_template.replace('{parameter_directory_code}', self.parameter_directory_code_map[parameter])
        data_grid_file = data_grid_file.replace('{year}{postfix}', year_str)
        data_grid_file = data_grid_file.replace('{parameter_file_code}', self.parameter_file_code_map[parameter])
        data_grid_file = data_grid_file.replace('{month_code}', self.month_codes[month_index])
        return { 'data_file' : data_grid_file, 'present' : path.exists(path.join(self.climate_data_directory['path'], data_grid_file)) }

    # convert data download interval label to AD year interval
    def convertDataDownloadIntervalLabelToAD(self, interval_label) :
        from_year = int(interval_label.split('-')[0][:-2])
        from_postfix = interval_label.split('-')[0][-2:]
        from_year_ad = from_year
        if from_postfix == 'BP' :
            from_year_ad = 1950 - from_year
        until_year = int(interval_label.split('-')[1][:-2])
        until_postfix = interval_label.split('-')[1][-2:]
        until_year_ad = until_year
        if until_postfix == 'BP' :
            until_year_ad = 1950 - until_year
        return { 'from_year_ad' : from_year_ad, 'until_year_ad' : until_year_ad }

    # get generation status
    def getGenerationStatus(self) :
        return self.generation_status

    # Set region mask directory
    def setRegionMaskDirectory(self, path) :
        self.region_mask_directory = self.splitPath(path)

    # Set bias correction directory
    def setBiasCorrectionDirectory(self, path) :
        self.bias_correction_directory = self.splitPath(path)

    # Set File Generation directory
    def setFileGenerationDirectory(self, path) :
        self.file_generation_directory = self.splitPath(path)

    # Get File Generation directory path
    def getFileGenerationDirectoryPath(self) :
        return self.file_generation_directory['path']

    # Get File Generation directory name
    def getFileGenerationDirectoryName(self) :
        return self.file_generation_directory['name']

    # Get File Generation directory root
    def getFileGenerationDirectoryRoot(self) :
        return self.file_generation_directory['directory']

    # Set generation status
    def setGenerationStatus(self, status) :
        self.generation_status = status

    # Set generation progress value
    def setGenerationProgressValue(self, progress_value) :
        self.generation_progress_value = progress_value

    # Method splits path into a dictionary containing path, name (end), and (parent) directory
    def splitPath(self, full_path) :
        path_dictionary = { 'path' : full_path }
        split_path = path.split(full_path)
        path_dictionary['directory'] = split_path[0]
        path_dictionary['name'] = split_path[1]
        return path_dictionary

    # Method recursively creates a new directory path (eliminates duplicate existing subdirectories)
    def createDirectoryPath(self, dir_path) :
        root_dir_path = self.splitPath(dir_path)['directory']
        new_dir_name = self.splitPath(dir_path)['name']
        if not(path.exists(root_dir_path)) :
            root_dir_path = self.createDirectoryPath(root_dir_path)
        if new_dir_name == self.splitPath(root_dir_path)['name'] :
            return root_dir_path
        else :
            new_path = path.join(root_dir_path, new_dir_name)
            mkdir(new_path)
            return new_path
            
    # Method generates a time-stamped directory within the current file generation directory
    def generateTimestampedFileGenerationDirectory(self) :

        # Create directory name from date and time
        timestamped_dir_name = strftime("%d%b%Y_%I%M%p_%S", localtime()) + '.' + string.split('%.3f' % time(), '.')[1] + 's'

        # Create the directory within the current file generation directory
        new_file_generation_path = self.createDirectoryPath(path.join(self.file_generation_directory['path'], timestamped_dir_name))
        self.setFileGenerationDirectory(new_file_generation_path)

        # Add timestamped directory name to file generation directory structure
        self.file_generation_directory['timestamped'] = timestamped_dir_name

        return self.file_generation_directory

    # Is the File Generation directory timestamped?
    def fileGenerationDirectoryIsTimestamped(self) :
        return self.file_generation_directory.has_key('timestamped')

    # Method resets the file generation directory to its parent when a time-stamped directory was utilised
    def resetTimestampedFileGenerationDirectory(self) :
        if self.file_generation_directory.has_key('timestamped') :
            self.file_generation_directory = self.splitPath(self.file_generation_directory['directory'])

    # Method checks if the file generation directory is empty
    def fileFenerationDirectoryIsEmpty(self) :
        if listdir(self.file_generation_directory['path']) :
            return False
        else : # no contents
            return True

    # Method loads region mask
    def loadRegionMask(self, region_code, time_dependent=False) :
        if time_dependent :
            region_mask_dict = {}
            for year in np.arange(0, 21001, 100) :
                mask_file = path.join(self.region_mask_directory['path'], region_code, (region_code.split('-')[0] + '-' + str(year)+ 'BP.msk'))
                region_mask_dict[year] = np.genfromtxt(mask_file, delimiter=1)
            return region_mask_dict
        else :
            mask_file = path.join(self.region_mask_directory['path'], (region_code + '.msk'))
            return np.genfromtxt(mask_file, delimiter=1)

    # Method finds nearest time dependent region mask year
    def nearestTimeDependentRegionMaskYear(self, year) :
        if year > 21000 :
            return 21000
        else :
            return int(round((year >= 0)*year/100.0,0)*100)
        
    # Method generates requested parameter data
    def generateParameterData(self,
                              parameter_group_code,
                              parameter_code,
                              period_ad_from,
                              period_ad_until,
                              delta_ref_period_ad=None,
                              delta_as_percent=False,
                              interval_step=10,
                              interval_size=10,
                              month_indices=range(12),
                              region_mask=1,
                              generate_grids=True,
                              all_months=False,
                              correct_bias=False) :

        # Expand region mask when not a grid (=1 in tests)
        if type(region_mask) != np.ndarray and type(region_mask) != dict :
            region_mask = np.zeros((self.grid_height, self.grid_width)) + region_mask

        # Update current NetCDF data file availability
        self.getCurrentNetCdfDataIntervals()
        
        # Setup proxy for web-based climate data when required
        self.setupClimateDataProxy()

        # Reset bias correction data cache
        self.cached_bias_correction_data_grids = {}

        # Load non-gridded data when required
        if self.non_gridded_parameter_files.has_key(parameter_group_code) and self.non_gridded_parameter_files[parameter_group_code].has_key(parameter_code) :
            if delta_ref_period_ad :
                delta_interval = { 'ad_from' : (delta_ref_period_ad - interval_size/2), 'ad_until' : (delta_ref_period_ad + interval_size/2 - int(not bool(interval_size%2))) }
            else :
                delta_interval = None
            self.loadNongriddedDataFrame(parameter_group_code=parameter_group_code,
                                         parameter_code=parameter_code,
                                         period_ad_from=(period_ad_from - interval_size/2),
                                         period_ad_until=(period_ad_until + interval_size/2 - int(not bool(interval_size%2))),
                                         delta_interval=delta_interval)

        # Collect data
        parameter_data = []
        intervals = range(period_ad_from, period_ad_until+1, interval_step)
        if all_months :
            for month_index in range(12) :
                parameter_data_month = []
                for i, interval_from in enumerate(intervals) :
                    parameter_data_month.append(self.generateParameterDataInterval(parameter_group_code=parameter_group_code,
                                                                                   parameter_code=parameter_code,
                                                                                   interval_ad_from=(interval_from - interval_size/2),
                                                                                   interval_ad_until=(interval_from + interval_size/2 - int(not bool(interval_size%2))),
                                                                                   month_indices=[month_index],
                                                                                   correct_bias=correct_bias))
                    if self.application_gui != None :
                        self.application_gui.generation_status_bar['value'] += interval_size
                        self.application_gui.update_idletasks()
                        #print 'generateParameterData, all_months:', all_months, self.application_gui.generation_status_bar['value'], time()
                parameter_data.append(parameter_data_month)
        else :
            for i, interval_from in enumerate(intervals) :
                parameter_data.append(self.generateParameterDataInterval(parameter_group_code=parameter_group_code,
                                                                         parameter_code=parameter_code,
                                                                         interval_ad_from=(interval_from - interval_size/2),
                                                                         interval_ad_until=(interval_from + interval_size/2 - int(not bool(interval_size%2))),
                                                                         month_indices=month_indices,
                                                                         correct_bias=correct_bias))
                if self.application_gui != None :
                    self.application_gui.generation_status_bar['value'] += interval_size * len(month_indices)
                    self.application_gui.update_idletasks()
                    #print 'generateParameterData, all_months:', all_months, self.application_gui.generation_status_bar['value'], time()

        # Calculate delta values when required
        if delta_ref_period_ad :
            if all_months :
                for month_index in range(12) :
                    delta_ref_data = self.generateParameterDataInterval(parameter_group_code=parameter_group_code,
                                                                        parameter_code=parameter_code,
                                                                        interval_ad_from=(delta_ref_period_ad - interval_size/2),
                                                                        interval_ad_until=(delta_ref_period_ad + interval_size/2 - int(not bool(interval_size%2))),
                                                                        month_indices=[month_index],
                                                                        correct_bias=correct_bias)
                    parameter_data[month_index] = self.calculateDeltaValues(parameter_data[month_index], delta_ref_data, delta_as_percent=delta_as_percent, grid_data=self.parameterDataIsGridded(parameter_group_code, parameter_code))
            else :
                delta_ref_data = self.generateParameterDataInterval(parameter_group_code=parameter_group_code,
                                                                    parameter_code=parameter_code,
                                                                    interval_ad_from=(delta_ref_period_ad - interval_size/2),
                                                                    interval_ad_until=(delta_ref_period_ad + interval_size/2 - int(not bool(interval_size%2))),
                                                                    month_indices=month_indices,
                                                                    correct_bias=correct_bias)
                parameter_data = self.calculateDeltaValues(parameter_data, delta_ref_data, delta_as_percent=delta_as_percent, grid_data=self.parameterDataIsGridded(parameter_group_code, parameter_code))

        # Return statistics for series data requests when gridded
        if self.parameterDataIsGridded(parameter_group_code, parameter_code) and not generate_grids :
            if all_months :
                grid_statistics = []
                for month_index in range(12) :
                    region_masks = []
                    for i, interval_from in enumerate(intervals) :
                        if type(region_mask) == dict :
                            region_masks.append(region_mask[self.nearestTimeDependentRegionMaskYear(1950 - interval_from)])
                        else :
                            region_masks.append(region_mask)
                    grid_statistics.append(self.calculateGridRegionStatistics(parameter_data[month_index], region_masks))
                return grid_statistics
            else :
                region_masks = []
                for i, interval_from in enumerate(intervals) :
                    if type(region_mask) == dict :
                        region_masks.append(region_mask[self.nearestTimeDependentRegionMaskYear(1950 - interval_from)])
                    else :
                        region_masks.append(region_mask)
                return self.calculateGridRegionStatistics(parameter_data, region_masks)

        else : # Return data list
            return parameter_data

    # Method generates requested parameter data for a single interval
    def generateParameterDataInterval(self,
                                      parameter_group_code,
                                      parameter_code,
                                      interval_ad_from,
                                      interval_ad_until,
                                      month_indices,
                                      correct_bias=False) :
        #print 'generateParameterDataInterval from', interval_ad_from, 'until', interval_ad_until

        # Resolve parameter calculations if present
        if self.parameter_calculation.has_key(parameter_group_code) and self.parameter_calculation[parameter_group_code].has_key(parameter_code) :
            parameter_calculation = self.parameter_calculation[parameter_group_code][parameter_code]
            calculate_key = ''

            # Handle non-gridded parameter files differently
            parameter_data_is_gridded = self.parameterDataIsGridded(parameter_group_code, parameter_code)

            # Aggregation parameter
            if parameter_data_is_gridded :
                sum_for_average = np.zeros((self.grid_height, self.grid_width))
            else :
                sum_for_average = 0

            # Calculate for each year in interval
            for year_ad in range(interval_ad_from, interval_ad_until+1) :

                # Collect/calculate any averages and seasonality standard deviations and coefficients of variation
                if 'average' in parameter_calculation['calculate'].keys() :
                    calculate_key = 'average'
                elif 'stdev_seasonality' in parameter_calculation['calculate'].keys() :
                    calculate_key = 'stdev_seasonality'
                elif 'coeff_var_seasonality' in parameter_calculation['calculate'].keys() :
                    calculate_key = 'coeff_var_seasonality'
                if calculate_key == 'average' or calculate_key == 'stdev_seasonality' or calculate_key == 'coeff_var_seasonality' :
                    for parameter in parameter_calculation['parameters'] :
                        if parameter_calculation['calculate'][calculate_key].find(parameter) > -1 :
                            if parameter_data_is_gridded :
                                grid_values = self.loadClimateDataGrids(parameter, year_ad, month_indices, correct_bias)
                                exec(parameter + ' = grid_values')
                            else :
                                data_values = self.loadNongriddedClimateData(parameter_group_code, parameter_code, year_ad, month_indices)
                                exec(parameter + ' = data_values')
                    calculation_values = eval(parameter_calculation['calculate'][calculate_key])
                    average = calculation_values.mean(0)
                    stdev_seasonality = calculation_values.std(0)
                    coeff_var_seasonality = stdev_seasonality/average

                # Calculate any annual ranges
                if 'annual_range' in parameter_calculation['calculate'].keys() :
                    calculate_key = 'annual_range'
                    for parameter in parameter_calculation['parameters'] :
                        if parameter_calculation['calculate'][calculate_key].find(parameter) > -1 :
                            if parameter_data_is_gridded :
                                grid_values = self.loadClimateDataGrids(parameter, year_ad, range(12), correct_bias)
                                exec(parameter + ' = grid_values')
                            else :
                                data_values = self.loadNongriddedClimateData(parameter_group_code, parameter_code, year_ad, range(12))
                                exec(parameter + ' = data_values')
                    calculation_values = eval(parameter_calculation['calculate'][calculate_key])
                    annual_range = calculation_values.max(0) - calculation_values.min(0)

                # Aggregate calculated value for the year
                sum_for_average += eval(parameter_calculation['return'])

            # Calculate average across interval
            parameter_data_interval = sum_for_average/(interval_ad_until - interval_ad_from + 1)

            # Return parameter data for interval
            return parameter_data_interval

    # Method determines if the parameter data gridded?
    def parameterDataIsGridded(self, parameter_group_code, parameter_code) :
        return not (self.non_gridded_parameter_files.has_key(parameter_group_code) and self.non_gridded_parameter_files[parameter_group_code].has_key(parameter_code))

    # Method loads non-gridded climate data into a panda data frame (avoids repeated loads)
    def loadNongriddedDataFrame(self, parameter_group_code, parameter_code, period_ad_from, period_ad_until, delta_interval=None) : # TODO: load from URL

        # Load full data file
        data_file = self.non_gridded_parameter_files[parameter_group_code][parameter_code]
        if data_file :
            if self.climate_data_source == 'url' :
                data_file_url = self.resolveClimateDataUrl(parameter_code, data_file)
                try :
                    data_file = url.urlopen(data_file_url)
                    data_frame = pd.read_csv(data_file, delim_whitespace=True)
                except Exception, e :
                    exception_message = 'Could not open ' + parameter_code.replace('_',' ').title() + ' data. Expected climate data file at: \n' + data_file_url + '\n' + str(e)
                    raise Exception(exception_message)
            else :
                data_file = path.join(self.climate_data_directory['path'], data_file)
                if path.exists(data_file) :
                    data_frame = pd.read_csv(data_file, delim_whitespace=True)
                else :
                    #print 'TODO: handle missing climate data for :', parameter_group_code, parameter_code
                    split_path = self.splitPath(data_file)
                    exception_message = 'Could not find ' + parameter_code.replace('_',' ').title() + ' data. Expected climate data file: \n' + data_file
                    raise Exception(exception_message)
        else :
            raise Exception('The data location for ' + parameter_code.upper() + ' has not been defined yet')            

        # Reduce to required years + previous year (BP and AD) to make further processing quicker
        year_col = self.non_gridded_parameter_file_columns[parameter_group_code][parameter_code]['year']
        postfix_col = self.non_gridded_parameter_file_columns[parameter_group_code][parameter_code]['postfix']
        frame_keys = (data_frame[postfix_col] == 'EMPTY')
        ad_years = []
        bp_years = []
        for ad_year in range(period_ad_from-1, period_ad_until+1) :
            if ad_year > 1950 :
                ad_years.append(ad_year)
            else :
                bp_years.append(1950 - ad_year)
        if ad_years :
            frame_keys = frame_keys | ((data_frame[postfix_col] == 'AD') & (data_frame[year_col] >= min(ad_years)) & (data_frame[year_col] <= max(ad_years)))
        if bp_years :
            frame_keys = frame_keys | ((data_frame[postfix_col] == 'BP') & (data_frame[year_col] >= min(bp_years)) & (data_frame[year_col] <= max(bp_years)))
        if delta_interval :
            ad_delta_years = []
            bp_delta_years = []
            for ad_year in range(delta_interval['ad_from']-1, delta_interval['ad_until']+1) :
                if ad_year > 1950 :
                    ad_delta_years.append(ad_year)
                else :
                    bp_delta_years.append(1950 - ad_year)
            if ad_delta_years :
                frame_keys = frame_keys | ((data_frame[postfix_col] == 'AD') & (data_frame[year_col] >= min(ad_delta_years)) & (data_frame[year_col] <= max(ad_delta_years)))
            if bp_delta_years :
                frame_keys = frame_keys | ((data_frame[postfix_col] == 'BP') & (data_frame[year_col] >= min(bp_delta_years)) & (data_frame[year_col] <= max(bp_delta_years)))

        self.nongridded_data_frame = data_frame[frame_keys]

    # Method loads non-gridded climate data for the selected months for a given year
    def loadNongriddedClimateData(self, parameter_group, parameter, year_ad, month_indices) :

        # Load from preloaded file data if present
        if type(self.nongridded_data_frame) == pd.DataFrame :
            data_frame = self.nongridded_data_frame

            # Rearrange month indices if they cross years
            month_indices = self.rearrangeMonthIndices(month_indices)

            # Extract the data
            climate_data = []
            postfix_col = self.non_gridded_parameter_file_columns[parameter_group][parameter]['postfix']
            year_col = self.non_gridded_parameter_file_columns[parameter_group][parameter]['year']
            month_code_col = self.non_gridded_parameter_file_columns[parameter_group][parameter]['month_code']
            data_col = self.non_gridded_parameter_file_columns[parameter_group][parameter]['data']
            for i in month_indices :
                if i > month_indices[-1] : # crosses into the previous year
                    year = year_ad-1
                else :
                    year = year_ad
                postfix = 'AD'
                if year <= 1950 :
                    postfix = 'BP'
                    year = 1950 - year
                month_code = self.month_codes[i]
                data_value = data_frame[(data_frame[postfix_col] == postfix) & (data_frame[year_col] == year) & (data_frame[month_code_col] == month_code)][data_col].get_values()[0]
                if data_value :
                    climate_data.append(data_value)
            #print 'loadNongriddedClimateData', year_ad, month_indices, np.array(climate_data)
            return np.array(climate_data)
        else :
            return np.array([])

    # Method loads climate data for the selected months for a given year
    def loadClimateDataGrids(self, parameter, year_ad, month_indices, correct_bias=False) :

        # Re-order month indices when months cross years
        month_indices = self.rearrangeMonthIndices(month_indices)

        # Load climate data grids for each month
        data_grids = []
        for i in month_indices :
            if i > month_indices[-1] : # crosses into the previous year
                data_grid = self.loadClimateDataGrid(parameter, year_ad-1, i)
            else : 
                data_grid = self.loadClimateDataGrid(parameter, year_ad, i)
            if data_grid != None :
                data_grids.append(data_grid)
        climate_data_grids = np.array(data_grids)

        # Correct bias when required
        if correct_bias :
            bias_correction_data_grids = self.loadBiasCorrectionDataGrids(parameter, month_indices)
            if self.pre_bias_correction_bounds.has_key(parameter) :
                for data_grids, bounds in self.pre_bias_correction_bounds[parameter].items() :
                    for bound_type, bound_value in bounds.items() :
                        exec(data_grids + ' = self.applyBoundToDataGrid(' + data_grids + ', bound_type, bound_value)')
            return eval(self.bias_correction_calculation[parameter])
        else :
            return climate_data_grids

    # Method applies a bound to a data grid
    def applyBoundToDataGrid(self, data_grids, bound_type, bound_value) :
        if bound_type == 'lower' :
            return (data_grids >= bound_value)*data_grids + (data_grids < bound_value)*bound_value
        elif bound_type == 'upper' :
            return (data_grids <= bound_value)*data_grids + (data_grids > bound_value)*bound_value

    # Method loads climate data for a selected month of a given year
    def loadClimateDataGrid(self, parameter, year_ad, month_index) : # TODO: load from URL
        #print 'loadClimateDataGrid:', parameter, year_ad, month_index

        # Year and postfix
        postfix = 'AD'
        year = year_ad
        if year_ad <= 1950 :
            postfix = 'BP'
            year = 1950 - year_ad
        year_str = str(year) + postfix

        # Original code handles individual local or networked raw files
        if (self.climate_data_source == 'local' and not self.use_netCdf_data) or self.climate_data_source == 'url' :
            if self.parameter_directory_code_map.has_key(parameter) :
                data_file = self.climate_data_file_template.replace('{parameter_directory_code}', self.parameter_directory_code_map[parameter])
                data_file = data_file.replace('{year}', str(year))
                data_file = data_file.replace('{postfix}', postfix)
                data_file = data_file.replace('{parameter_file_code}', self.parameter_file_code_map[parameter])
                data_file = data_file.replace('{month_code}', self.month_codes[month_index])
                if self.climate_data_source == 'url' :
                    data_file_url = self.resolveClimateDataUrl(parameter, data_file)
                    try :
                        data_file = url.urlopen(data_file_url)
                        return np.genfromtxt(data_file)
                    except Exception, e :
                        exception_message = 'Could not open ' + parameter.replace('_', ' ').title() + ' data for ' + self.month_names[month_index] + ' ' + str(year) + postfix + '. Expected climate data file at: \n' + data_file_url + '\n' + str(e)
                        raise Exception(exception_message)
                else :
                    data_file = path.join(self.climate_data_directory['path'], data_file)
                    if path.exists(data_file) :
                        return np.genfromtxt(data_file)
                    else :
                        #print 'TODO: handle missing climate data:', path.join(self.splitPath(split_path['directory'])['name'], split_path['name'])
                        exception_message = 'Could not find ' + parameter.replace('_', ' ').title() + ' data for ' + self.month_names[month_index] + ' ' + str(year) + postfix + '. Expected climate data file: \n' + data_file
                        raise Exception(exception_message)
            else :
                raise Exception('The data location for ' + parameter.replace('_', ' ').title() + ' has not been defined yet')

        # Utilise local NetCDF files when present
        elif (self.climate_data_source == 'local' and self.use_netCdf_data) :

            # Check NetCDF cache for data first
            if self.cached_netCdf_data.has_key(parameter) :

                # Get root group DataSet object
                rootgrp = self.cached_netCdf_data[parameter]['rootgrp']

                # Check current subgroup
                sub_interval_ad = self.convertDataIntervalLabelToAD(self.cached_netCdf_data[parameter]['sub_interval_str'])
                if sub_interval_ad['from_year_ad'] <= year_ad <= sub_interval_ad['until_year_ad'] :
                    subgroup = rootgrp.groups[self.cached_netCdf_data[parameter]['sub_interval_str']]
                    return subgroup.variables[year_str][month_index]

##                # Check current rootgrp
##                for sub_interval_str in rootgrp.groups.keys() :
##                    sub_interval_ad = self.convertDataIntervalLabelToAD(sub_interval_str)
##                    if sub_interval_ad['from_year_ad'] <= year_ad <= sub_interval_ad['until_year_ad'] :
##                        self.cached_netCdf_data[parameter]['sub_interval_str'] = sub_interval_str
##                        return rootgrp.groups[sub_interval_str].variables[year_str][month_index]

                # Not found - clear cache for parameter
                rootgrp.close()
                self.cached_netCdf_data.pop(parameter)

            # Check available NetCDF data files
            for root_interval_str in self.current_netCdf_data_intervals[parameter] :
                root_interval_ad = self.convertDataIntervalLabelToAD(root_interval_str)
                root_from_year_ad = max((root_interval_ad['from_year_ad'] - self.download_data_window), (1950 - 22000))
                root_until_year_ad = min((root_interval_ad['until_year_ad'] + self.download_data_window), 1989)
                if root_from_year_ad <= year_ad <= root_until_year_ad :
                    try :
                        rootgrp = Dataset(path.join(self.climate_data_directory['path'], (parameter+'-'+root_interval_str+'.nc')), 'r')
                        for sub_interval_str in rootgrp.groups.keys() :
                            sub_interval_ad = self.convertDataIntervalLabelToAD(sub_interval_str)
                            if sub_interval_ad['from_year_ad'] <= year_ad <= sub_interval_ad['until_year_ad'] :
                                self.cached_netCdf_data[parameter] = { 'rootgrp' : rootgrp, 'root_interval_str' : root_interval_str, 'sub_interval_str' : sub_interval_str  }
                                return rootgrp.groups[sub_interval_str].variables[year_str][month_index]
                    except Exception, e :
                        exception_message = 'Could not open NetCDF data file: ' + path.join(self.climate_data_directory['path'], (parameter+'-'+root_interval_str+'.nc')) + '\n' + str(e)
                        raise Exception(exception_message)

            # Still not found
            raise Exception('Could not find NetCDF data for '+ parameter.replace('_', ' ').title() + ' data for ' + self.month_names[month_index] + ' ' + str(year) + postfix + '.\n')

    # Method loads bias correction data for the selected months
    def loadBiasCorrectionDataGrids(self, parameter, month_indices) :
        month_indices = self.rearrangeMonthIndices(month_indices)
        data_grids = []
        for i in month_indices :
            data_grids.append(self.loadBiasCorrectionDataGrid(parameter, i))
        return np.array(data_grids)

    # Method loads bias correction data for a selected month
    def loadBiasCorrectionDataGrid(self, parameter, month_index) :
        if self.cached_bias_correction_data_grids.has_key(parameter) and self.cached_bias_correction_data_grids[parameter].has_key(month_index) :
            return self.cached_bias_correction_data_grids[parameter][month_index]
        elif self.bias_correction_directory_code_map.has_key(parameter) :
            data_file = self.bias_correction_file_template.replace('{bias_correction_directory_code}', self.bias_correction_directory_code_map[parameter])
            data_file = data_file.replace('{parameter_file_code}', self.parameter_file_code_map[parameter])
            data_file = data_file.replace('{month_code}', self.month_codes[month_index])
            data_file = data_file.replace('{bias_correction_file_code}', self.bias_correction_file_code_map[parameter])
            data_file = path.join(self.bias_correction_directory['path'], data_file)
            if path.exists(data_file) :
                data_grid = np.genfromtxt(data_file)
                if not self.cached_bias_correction_data_grids.has_key(parameter) :
                    self.cached_bias_correction_data_grids[parameter] = {}
                self.cached_bias_correction_data_grids[parameter][month_index] = data_grid
                return data_grid
            else :
                #print 'TODO: handle missing climate data:', path.join(self.splitPath(split_path['directory'])['name'], split_path['name'])
                exception_message = 'Could not find ' + parameter.replace('_', ' ').title() + ' bias correction data for ' + self.month_names[month_index] + '. Expected bias correction data file: \n' + data_file
                raise Exception(exception_message)
        else :
            raise Exception('The bias correction data location for ' + parameter.replace('_', ' ').title() + ' has not been defined yet')

    # Method rearrange month indices if they cross years
    def rearrangeMonthIndices(self, month_indices) :
        month_indices = month_indices[:]
        month_indices.sort()
        if len(month_indices) < 12 and month_indices.count(0) and month_indices.count(11) : # months cross years
            this_year_indices = [month_indices.pop()]
            while month_indices[:].pop() == (this_year_indices[0] - 1) :
                this_year_indices.insert(0, month_indices.pop())
            this_year_indices.extend(month_indices)
            return this_year_indices
        else :
            return month_indices

    # Method calculates delta values
    def calculateDeltaValues(self, parameter_data, delta_ref_data, delta_as_percent=False, grid_data=True) :

        # Handle gridded data
        if grid_data :
            delta_data = []
            for i, data_grid in enumerate(parameter_data) :
                if delta_as_percent :
                    delta_grid = 100*(data_grid - delta_ref_data)/delta_ref_data
                    delta_grid[np.isnan(delta_grid).nonzero()] = 0.0
                    delta_grid[(np.isinf(delta_grid)*(delta_grid < 0)).nonzero()] = delta_grid[np.isfinite(delta_grid).nonzero()].min()
                    delta_grid[(np.isinf(delta_grid)*(delta_grid > 0)).nonzero()] = delta_grid[np.isfinite(delta_grid).nonzero()].max()
                    delta_data.append(delta_grid)
                else :
                    delta_data.append(data_grid - delta_ref_data)

        # Handle non-gridded data
        else :
            data_array = np.array(parameter_data)
            if delta_as_percent :
                delta_array = 100*(data_array - delta_ref_data)/delta_ref_data
                delta_array[np.isnan(delta_array).nonzero()] = 0.0
                delta_array[(np.isinf(delta_array)*(delta_array < 0)).nonzero()] = -100.0
                delta_array[(np.isinf(delta_array)*(delta_array > 0)).nonzero()] = 100.0
                delta_data = delta_array.tolist()
            else :
                delta_data = (data_array - delta_ref_data).tolist()

        return delta_data

    # Method calculates grid region statistics
    def calculateGridRegionStatistics(self, grid_data_list, region_mask_list) :
        grid_statistics = { 'minimum' : [], 'percentile_5th' : [], 'percentile_25th' : [], 'percentile_50th' : [], 'percentile_75th' : [], 'percentile_95th' : [], 'maximum' : [],
                            'grid_mean' : [], 'grid_stdev' : [], 'area_mean' : [], 'area_stdev' : [] }
        for i, grid_data in enumerate(grid_data_list) :
            finite_grid_data = np.ma.masked_array(grid_data, mask=((np.isfinite(grid_data) - 1)*-1))
            region_data = finite_grid_data[region_mask_list[i].nonzero()]
            grid_statistics['minimum'].append(region_data.min())
            grid_statistics['percentile_5th'].append(np.percentile(region_data, 5))
            grid_statistics['percentile_25th'].append(np.percentile(region_data, 25))
            grid_statistics['percentile_50th'].append(np.percentile(region_data, 50))
            grid_statistics['percentile_75th'].append(np.percentile(region_data, 75))
            grid_statistics['percentile_95th'].append(np.percentile(region_data, 95))
            grid_statistics['maximum'].append(region_data.max())
            grid_statistics['grid_mean'].append(region_data.mean())
            grid_statistics['grid_stdev'].append(region_data.std())
            area_weighted_region_mask = region_mask_list[i] * np.meshgrid(np.ones(144), np.cos(np.arange(88.75,-88.751,-2.5)*np.pi/180.0))[1]
            area_mean = (finite_grid_data*area_weighted_region_mask).sum()/area_weighted_region_mask.sum()
            grid_statistics['area_mean'].append(area_mean)
            grid_statistics['area_stdev'].append(np.sqrt(((area_mean-finite_grid_data)**2*area_weighted_region_mask).sum()/area_weighted_region_mask.sum()))
        return grid_statistics

    # Method generates a grid data file
    def generateGridDataFile(self, masked_data, file_type, year_label, description=None, data_units=None) :
        masked_data[((np.isfinite(masked_data)-1)*-1).nonzero()] = np.nan
        data_frame = pd.DataFrame(masked_data)
        if file_type == 'csv' :
            output_file_path = path.join(self.file_generation_directory['path'], ('grid_data_'+year_label+'.csv'))
            data_frame.to_csv(output_file_path, header=False, index=False, na_rep='-9999', float_format='%.3f')
        elif file_type == 'ascii' :
            output_file_path = path.join(self.file_generation_directory['path'], ('grid_data_'+year_label+'.txt'))
            f = open(output_file_path, 'w')
            f.write(data_frame.to_string(header=False, index=False, nanRep='-9999', na_rep='-9999', float_format=(lambda f: '%.3f'%f)).replace('  nan','-9999'))
            f.close()
        elif file_type == 'esri_ascii' :
            header = 'ncols 144\n' + 'nrows 72\n' + 'xllcorner -180\n' + 'yllcorner -90\n' + 'cellsize 2.5\n' + 'nodata_value -9999\n'
            string_buffer = StringIO()
            data_frame.to_csv(string_buffer, sep=' ', header=False, index=False, na_rep='-9999', float_format='%.3f')
            output_file_path = path.join(self.file_generation_directory['path'], ('grid_data_'+year_label+'.asc'))
            f = open(output_file_path, 'w')
            f.write(header)
            f.write(string_buffer.getvalue())
            f.close()
        elif file_type == 'netcdf' :
            output_file_path = path.join(self.file_generation_directory['path'], ('grid_data_'+year_label+'.nc'))
            rootgrp = Dataset(output_file_path, 'w')
            rootgrp.description = str(description)
            lat = rootgrp.createDimension('lat', 72)
            lon = rootgrp.createDimension('lon', 144)
            latitudes = rootgrp.createVariable('latitudes','f8',('lat',))
            longitudes = rootgrp.createVariable('longitudes','f8',('lon',))
            data = rootgrp.createVariable('data','f8',('lat','lon',))
            latitudes.units = 'degrees north'
            longitudes.units = 'degrees east'
            data.units = str(data_units)
            latitudes[:] = np.arange(88.75,-88.751,-2.5)
            longitudes[:] = np.arange(-178.75,178.751,2.5)
            data[:,:] = masked_data 
            rootgrp.close()

    # Method generates a series data file
    def generateSeriesDataFile(self, data_frame, file_type, month='', description=None, data_units=None) :
        if month :
            month = '_' + month
        if file_type == 'csv' :
            output_file_path = path.join(self.file_generation_directory['path'], ('series_data'+month+'.csv'))
            data_frame.to_csv(output_file_path, index=False)
        elif file_type == 'ascii' :
            output_file_path = path.join(self.file_generation_directory['path'], ('series_data'+month+'.txt'))
            f = open(output_file_path, 'w')
            f.write(data_frame.to_string(index=False))
            f.close()
        elif file_type == 'netcdf' :
            output_file_path = path.join(self.file_generation_directory['path'], ('series_data'+month+'.nc'))
            rootgrp = Dataset(output_file_path, 'w')
            rootgrp.description = str(description)
            series = rootgrp.createDimension('series', data_frame.shape[0])
            variables = {}
            types = { 'int64' : 'i8', 'float64' : 'f8' }
            for i, key in enumerate(data_frame.keys()) :
                if types.has_key(str(data_frame[key].dtype)) :
                    data_type = types[str(data_frame[key].dtype)]
                else :
                    if i == 0 :
                        data_type = 'i8'
                    else :
                        data_type = 'f8'
                variables[key] = rootgrp.createVariable(key, data_type, ('series',))
                if i > 0 :
                    variables[key].units = str(data_units)
                variables[key][:] = data_frame[key].as_matrix()
            rootgrp.close()

    # Method generates a log entry time-stamp
    def generateLogEntryTimestamp(self) :
        return strftime("%d %b %Y %I:%M%p", localtime())

    def updateToolGenerationLogs(self, logfiles, full_log_entry, log_entry_addition=None) :
        for logfile in logfiles :
            if path.exists(logfile) :
                f = open(logfile, 'a')
                if log_entry_addition :
                    f.write(log_entry_addition)
                else :
                    f.write(full_log_entry)
            else :
                f = open(logfile, 'w')
                f.write(full_log_entry)
            f.close()

    # Method resolves the URL for a given data file.
    def resolveClimateDataUrl(self, parameter, data_file) :
        #print 'resolveClimateDataUrl', parameter, data_file
        data_file = data_file.replace('\\','/')
        return self.climate_data_url + data_file

    # Method generates a Word (docx) table from a data frame and title : discontinued use
    def generateWordTable(self, data_frame, title, ordered_columns=None, file_path=None) :
        if not ordered_columns :
            ordered_columns = data_frame
        if not file_path :
            file_path = path.join(self.file_generation_directory['path'], 'grid_plot_statistics.docx')
        document = Document()
        document.add_heading(title)
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        table = document.add_table(rows=(data_frame.shape[0]+1), cols=data_frame.shape[1])
        for i, column in enumerate(ordered_columns) :
            table.cell(0, i).text = column
            for row in range(data_frame.shape[0]) :
                table.cell(row+1, i).text = str(data_frame[column][row])
        document.save(file_path)

    # Method generates a Text (txt) table from a data frame and title : utilised for generating grid plot statistics tables (so far)
    def generateTextTable(self, data_frame, title, ordered_columns=None, file_path=None) :
        if not ordered_columns :
            ordered_columns = data_frame
        if not file_path :
            file_path = path.join(self.file_generation_directory['path'], 'grid_plot_statistics.txt')
        f = open(file_path, 'w')
        f.write(title + '\n' + '-'*len(title) + '\n' + data_frame.to_string(index=False, float_format=(lambda f: '%.3f'%f)))
        f.close()

    # Get Current NetCDF Data Intervals and use NetCDF data files if some are present
    def getCurrentNetCdfDataIntervals(self) :
        self.current_netCdf_data_intervals = {}
        for parameter in self.parameter_directory_code_map.keys() :
            self.current_netCdf_data_intervals[parameter] = []
            for data_interval_str in self.climate_data_download_intervals :
                expected_netCdf_file = (parameter+'-'+data_interval_str+'.nc')
                if path.exists(path.join(self.climate_data_directory['path'], expected_netCdf_file)) :
                    self.current_netCdf_data_intervals[parameter].append(data_interval_str)
                    self.use_netCdf_data = True
        return self.current_netCdf_data_intervals
                
    # Convert data interval label to AD year interval
    def convertDataIntervalLabelToAD(self, interval_label) :
        from_year = int(interval_label.split('-')[0][:-2])
        from_postfix = interval_label.split('-')[0][-2:]
        from_year_ad = from_year
        if from_postfix == 'BP' :
            from_year_ad = 1950 - from_year
        until_year = int(interval_label.split('-')[1][:-2])
        until_postfix = interval_label.split('-')[1][-2:]
        until_year_ad = until_year
        if until_postfix == 'BP' :
            until_year_ad = 1950 - until_year
        return { 'from_year_ad' : from_year_ad, 'until_year_ad' : until_year_ad }

    # Convert AD year interval to data interval label
    def convertAdIntervalToDataLabel(self, from_year_ad, until_year_ad) :
        if from_year_ad > 1950 :
            from_year_str = str(from_year_ad) + 'AD'
        else :
            from_year_str = str(1950 - from_year_ad) + 'BP'
        if until_year_ad > 1950 :
            until_year_str = str(until_year_ad) + 'AD'
        else :
            until_year_str = str(1950 - until_year_ad) + 'BP'
        return from_year_str + '-' + until_year_str

    # Method generates a NetCDF file for a climate data parameter for the specified interval 
    def generateNetCdfClimateData(self, parameter, from_year_ad, until_year_ad, min_year_ad=(1950-22000), max_year_ad=1989, zlib=True, decimals=None) :

        # Construct NetCDF file path
        data_interval_str = self.convertAdIntervalToDataLabel(from_year_ad, until_year_ad)
        [from_year_str, until_year_str] = data_interval_str.split('-')
        output_file_path = path.join(self.file_generation_directory['path'], (parameter+'-'+data_interval_str+'.nc'))
        print 'Generating NetCDF data file for ' + parameter.replace('_',' ').title() + ' ' + data_interval_str ###

        # Resolve data format. Assumes data in all files are the same as the first.
        first_grid_file = self.climate_data_file_template.replace('{parameter_directory_code}', self.parameter_directory_code_map[parameter])
        first_grid_file = first_grid_file.replace('{year}{postfix}', from_year_str)
        first_grid_file = first_grid_file.replace('{parameter_file_code}', self.parameter_file_code_map[parameter])
        first_grid_file = first_grid_file.replace('{month_code}', self.month_codes[0])
        first_line = open(path.join(self.climate_data_directory['path'], first_grid_file), 'r').readline()
        data_width = len(first_line.split('.')[1]) + 1
        if decimals :
            data_decimals = decimals
        else :
            data_decimals = len(first_line.split('.')[1].split(' '))

        # Construct NetCDF dataset file
        rootgrp = Dataset(output_file_path, 'w')
        rootgrp.Conventions = 'CF-1.6'
        rootgrp.title = parameter.replace('_',' ').title() + ' ' + from_year_str + ' - ' + until_year_str
        rootgrp.institution = 'Global Ecology Lab'
        rootgrp.source = 'https://github.com/GlobalEcologyLab/PaleoView'
        rootgrp.history = '[' + strftime("%Y-%m-%d %H:%M", localtime()) + ']' + 'Created netCDF4 zlib=True dataset.'
        rootgrp.description = parameter.replace('_',' ').title() + ' ' + from_year_str + ' - ' + until_year_str
        rootgrp.createDimension('single', 1)
        rootgrp.createDimension('month', 12)
        rootgrp.createDimension('lat', 72)
        rootgrp.createDimension('lon', 144)
        window = rootgrp.createVariable('window','i4',('single',))
        width = rootgrp.createVariable('width','i4',('single',))
        decimals = rootgrp.createVariable('decimals','i4',('single',))
        months = rootgrp.createVariable('months','i4',('month',), zlib=zlib)
        latitudes = rootgrp.createVariable('latitudes','f8',('lat',), zlib=zlib)
        longitudes = rootgrp.createVariable('longitudes','f8',('lon',), zlib=zlib)
        latitudes.units = 'degrees north'
        longitudes.units = 'degrees east'
        window[:] = np.array(self.download_data_window)
        width[:] = np.array(data_width)
        decimals[:] = np.array(data_decimals)
        months[:] = np.array(range(12)) + 1
        latitudes[:] = np.arange(88.75,-88.751,-2.5)
        longitudes[:] = np.arange(-178.75,178.751,2.5)

        # Add overlap to allow for interval window
        from_year_ad = max((from_year_ad - self.download_data_window), min_year_ad)
        until_year_ad = min((until_year_ad + self.download_data_window), max_year_ad)

        # Create subgoup intervals from start, 1000 year BP boundaries, and AD section
        group_from = [from_year_ad]
        group_until = []
        bp_1000_boundaries = (1950 - np.arange(((1950-from_year_ad)/1000*1000), (1950-until_year_ad)-1, -1000)).tolist()
        if from_year_ad in bp_1000_boundaries :
            bp_1000_boundaries.pop(0)
        group_from.extend(bp_1000_boundaries)
        if until_year_ad in bp_1000_boundaries :
            bp_1000_boundaries.pop()
        group_until.extend(bp_1000_boundaries)
        if from_year_ad < 1950 and until_year_ad > 1950 :
            group_from[-1] = 1951
        group_until.append(until_year_ad)
        subgroup_ad_intervals = np.array([group_from, group_until]).transpose()

        # Create subgroups and variables for each year's data within each subgroup
        for [subgroup_from_year_ad, subgroup_until_year_ad] in subgroup_ad_intervals :

            # Create subgrooup for interval
            subgroup_interval_str = self.convertAdIntervalToDataLabel(subgroup_from_year_ad, subgroup_until_year_ad)
            subgroup = rootgrp.createGroup(subgroup_interval_str)

            # Add a variable for each year within subgroups
            for year_ad in range(subgroup_from_year_ad, subgroup_until_year_ad+1) : 

                # Create year label 
                if year_ad > 1950 :
                    year_str = str(year_ad) + 'AD'
                else :
                    year_str = str(1950 - year_ad) + 'BP'

                # Load the year's climate data grids
                year_grids = []
                for i in range(12) :
                    year_grids.append(self.loadClimateDataGrid(parameter, year_ad, i))

                # Load grid data into subgroup variable using year label
                data = subgroup.createVariable(year_str,'f8',('month','lat','lon',), zlib=zlib, least_significant_digit=(data_decimals+1)) # one more decimal than formatted
                data.units = self.parameter_unit_string[parameter]
                data.long_name = parameter.replace('_',' ').title()
                data.standard_name = parameter
                data[:,:,:] = np.array(year_grids)

                if year_ad % 100 == 0 : ### 
                    print year_str, strftime("%Y-%m-%d %H:%M", localtime())

            # Re-open NetCDF file for every subgroup (avoids python crash)
            rootgrp.close()
            rootgrp = Dataset(output_file_path, 'a')

        # Close NetCDF file
        rootgrp.close()

    # Method downloads and unpacks a NetCDF file for a climate data parameter for the interval expressed as a label
    def downloadClimateDataInterval(self, parameter, interval_label, delimiter='', retain_netCdf_file=True, unpack=False) :
        interval = self.convertDataIntervalLabelToAD(interval_label)
        self.downloadClimateData(parameter, interval['from_year_ad'], interval['until_year_ad'], delimiter=delimiter, retain_netCdf_file=retain_netCdf_file, unpack=unpack)

    # Method downloads and unpacks a NetCDF file for a climate data parameter for the specified interval 
    def downloadClimateData(self, parameter, from_year_ad, until_year_ad, delimiter='', retain_netCdf_file=True, unpack=False) :

        # Construct NetCDF file name
        data_interval_str = self.convertAdIntervalToDataLabel(from_year_ad, until_year_ad)
        netCdf_file = (parameter+'-'+data_interval_str+'.nc')

        # Setup proxy for web-based climate data when required
        self.setupClimateDataProxy()

        # Download NetCDF file
        local_netCdf_path = path.join(self.climate_data_directory['path'], netCdf_file)
        try :
            check_connection = url.urlopen((self.climate_data_url + netCdf_file))
            check_connection.close()
            urllib.urlretrieve((self.climate_data_url + netCdf_file), local_netCdf_path, reporthook=self.netCdfDownloadProgress)
        except Exception, e :
            exception_message = 'Could not open ' + netCdf_file + '\nExpected climate data NetCDF file at: \n' + self.climate_data_url + '\n' + str(e)
            raise Exception(exception_message)
       
        # Unpack climate data files from NetCDF dataset file
        if unpack :
            rootgrp = Dataset(local_netCdf_path, 'r')
            width = rootgrp.variables['width'][0]
            decimals = rootgrp.variables['decimals'][0]
            if not(path.exists(path.join(self.climate_data_directory['path'], self.parameter_directory_code_map[parameter]))) :
                self.createDirectoryPath(path.join(self.climate_data_directory['path'], self.parameter_directory_code_map[parameter]))
            for subgroup_interval_str, subgroup in rootgrp.groups.items() :
                for year_str, year_variable in subgroup.variables.items() :
                    if self.parameter_directory_code_map.has_key(parameter) :
                        data_file_template = self.climate_data_file_template.replace('{parameter_directory_code}', self.parameter_directory_code_map[parameter])
                        data_file_template = data_file_template.replace('{year}{postfix}', year_str)
                        data_file_template = data_file_template.replace('{parameter_file_code}', self.parameter_file_code_map[parameter])
                        for month_index in range(12) :
                            data_file_path = path.join(self.climate_data_directory['path'], data_file_template.replace('{month_code}', self.month_codes[month_index]))
                            np.savetxt(data_file_path, subgroup.variables[year_str][month_index], fmt='%'+str(width)+'.'+str(decimals)+ 'f', delimiter=delimiter)
                # Re-open NetCDF file for every subgroup (avoids python crash)
                rootgrp.close()
                rootgrp = Dataset(local_netCdf_path, 'r')
            rootgrp.close()

        # Remove downloaded NetCDF file when not retained
        if not retain_netCdf_file :
            remove(local_netCdf_path)

    # Method tracks download progress and updates status bar
    def netCdfDownloadProgress(self, count, block_size, total_size) :

        if self.application_gui != None :

            # Set status maximum if not already set
            if self.application_gui.climate_data_download_status_bar['maximum'] == 1 :
                self.application_gui.climate_data_download_status_bar['maximum'] = total_size/block_size/1000
                self.application_gui.update_idletasks()

            # Set status value every 1000 blocks
            if not count % 1000 :
                self.application_gui.climate_data_download_status_bar['value'] += 1
                self.application_gui.update_idletasks()
